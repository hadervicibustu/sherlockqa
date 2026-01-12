from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

def create_bookmark_id(title):
    """Create a valid bookmark ID from a title."""
    # Remove special characters and replace spaces with underscores
    bookmark_id = re.sub(r'[^\w\s-]', '', title)
    bookmark_id = re.sub(r'\s+', '_', bookmark_id)
    return f"_Ref_{bookmark_id}"

def add_bookmark(paragraph, bookmark_name):
    """Add a bookmark to a paragraph."""
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    tag = run._r

    # Create bookmark start
    bookmark_start = OxmlElement('w:bookmarkStart')
    bookmark_start.set(qn('w:id'), '0')
    bookmark_start.set(qn('w:name'), bookmark_name)

    # Create bookmark end
    bookmark_end = OxmlElement('w:bookmarkEnd')
    bookmark_end.set(qn('w:id'), '0')

    # Insert bookmark
    tag.insert(0, bookmark_start)
    tag.append(bookmark_end)

def add_hyperlink(paragraph, text, bookmark_name, font_size=Pt(11), bold=False, indent=0):
    """Add an internal hyperlink to a bookmark."""
    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('w:anchor'), bookmark_name)

    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')

    # Set run properties
    rPr = OxmlElement('w:rPr')

    # Add color (blue for links)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Add font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)

    # Add bold if needed
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    # Set paragraph indentation
    if indent > 0:
        paragraph.paragraph_format.left_indent = Inches(indent)

    # Set line spacing
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(2)

def add_heading_with_bookmark(doc, text, level):
    """Add a heading with a bookmark."""
    heading = doc.add_heading(text, level=level)
    bookmark_name = create_bookmark_id(text)
    add_bookmark(heading, bookmark_name)
    return heading

def add_external_hyperlink(paragraph, url, text, font_size=Pt(11)):
    """Add an external hyperlink to a URL."""
    # Get the document part
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create the hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create a new run for the hyperlink text
    new_run = OxmlElement('w:r')

    # Set run properties
    rPr = OxmlElement('w:rPr')

    # Add color (blue for links)
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)

    # Add underline
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    # Add font size
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size.pt * 2)))
    rPr.append(szCs)

    new_run.append(rPr)

    # Add text
    text_elem = OxmlElement('w:t')
    text_elem.text = text
    new_run.append(text_elem)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

# Create a new Document
doc = Document()

# Add title
title = doc.add_heading('Software Requirements Specification', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add subtitle
subtitle = doc.add_paragraph('Ask Holmes Application')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add GitHub link
github_para = doc.add_paragraph()
github_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
github_run = github_para.add_run('GitHub: ')
github_run.font.size = Pt(11)
add_external_hyperlink(github_para, 'https://github.com/hadervicibustu/sherlockqa', 'https://github.com/hadervicibustu/sherlockqa', font_size=Pt(11))

# Add Live link
live_para = doc.add_paragraph()
live_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
live_run = live_para.add_run('Live: ')
live_run.font.size = Pt(11)
add_external_hyperlink(live_para, 'https://sherlockqa-oht2j.ondigitalocean.app/', 'https://sherlockqa-oht2j.ondigitalocean.app/', font_size=Pt(11))

# Add version info
version = doc.add_paragraph('Version 1.0')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add section break for title page (to allow vertical centering)
doc.add_section()

# Set vertical alignment to center for the first section (title page)
first_section = doc.sections[0]
sectPr = first_section._sectPr
vAlign = OxmlElement('w:vAlign')
vAlign.set(qn('w:val'), 'center')
sectPr.append(vAlign)

# Add page numbers to footer (skip first page/section)
# Get the second section (content pages)
second_section = doc.sections[1]

# Unlink footer from previous section so first page has no page number
second_section.footer.is_linked_to_previous = False

# Add page number to footer
footer = second_section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Create PAGE field for page number
run = footer_para.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')

instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'PAGE'

fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')

fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')

run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)

# ============================================
# TABLE OF CONTENTS
# ============================================
toc_heading = doc.add_heading('Table of Contents', level=1)
toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Define TOC entries: (level, title)
# Level 1 = main sections, Level 2 = subsections, Level 3 = sub-subsections
toc_entries = [
    (1, '1. Introduction'),
    (2, '1.1 Purpose'),
    (2, '1.2 Document Conventions'),
    (2, '1.3 Intended Audience'),
    (2, '1.4 Product Scope'),
    (2, '1.5 Product Overview'),
    (2, '1.6 References'),
    (2, '1.7 Definitions, Acronyms, and Abbreviations'),
    (3, '1.7.1 Acronyms'),
    (3, '1.7.2 Definitions'),
    (3, '1.7.3 Technical Terms'),
    (1, '2. Project Requirements'),
    (2, '2.1 Functional Requirements'),
    (3, '2.1.1 User Authentication Requirements'),
    (3, '2.1.2 Question Management Requirements'),
    (3, '2.1.3 Answer Management Requirements'),
    (3, '2.1.4 RAG Pipeline Requirements'),
    (3, '2.1.5 User Interface Requirements'),
    (3, '2.1.6 Data Management Requirements'),
    (2, '2.2 Non-Functional Requirements'),
    (3, '2.2.1 Performance Requirements'),
    (3, '2.2.2 Security Requirements'),
    (3, '2.2.3 Reliability Requirements'),
    (3, '2.2.4 Usability Requirements'),
    (3, '2.2.5 Accessibility Requirements'),
    (3, '2.2.6 Maintainability Requirements'),
    (3, '2.2.7 Scalability Requirements'),
    (3, '2.2.8 Compatibility Requirements'),
    (3, '2.2.9 Deployment Requirements'),
    (1, '3. Use Case Diagram'),
    (2, '3.1 Actors'),
    (2, '3.2 Use Cases by Package'),
    (3, '3.2.1 Authentication Package'),
    (3, '3.2.2 Question Management Package'),
    (3, '3.2.3 Answer Management Package'),
    (3, '3.2.4 RAG Pipeline Package'),
    (3, '3.2.5 Document Management Package'),
    (3, '3.2.6 User Interface Package'),
    (2, '3.3 Relationships Summary'),
    (2, '3.4 PlantUML Diagram Code'),
    (1, '4. Activity Diagram - Questions CRUD Operations'),
    (2, '4.1 Create Question'),
    (2, '4.2 Read Questions (View List)'),
    (2, '4.3 Read Question (View Details)'),
    (2, '4.4 Update Question'),
    (2, '4.5 Delete Question'),
    (2, '4.6 PlantUML Diagram Code'),
    (1, '5. Sequence Diagram - Questions CRUD Operations'),
    (2, '5.1 Participants'),
    (2, '5.2 Create Question Sequence'),
    (2, '5.3 Read Questions (List) Sequence'),
    (2, '5.4 Read Question (Details) Sequence'),
    (2, '5.5 Update Question Sequence'),
    (2, '5.6 Delete Question Sequence'),
    (2, '5.7 Sequence Diagram'),
    (1, '6. Class Diagram'),
    (2, '6.1 Overview'),
    (2, '6.2 Backend Models'),
    (2, '6.3 Backend Services'),
    (2, '6.4 Backend Utilities'),
    (2, '6.5 Backend API Routes'),
    (2, '6.6 Frontend Components'),
    (2, '6.7 Frontend Services and Context'),
    (2, '6.8 External Systems'),
    (2, '6.9 Key Relationships'),
    (2, '6.10 Class Diagram'),
    (1, '7. Entity Relationship Diagram'),
    (2, '7.1 Overview'),
    (2, '7.2 Entities'),
    (2, '7.3 Relationships'),
    (2, '7.4 Indexes and Performance'),
    (2, '7.5 Entity Relationship Diagram'),
    (1, '8. Design Patterns'),
    (2, '8.1 Backend Design Patterns'),
    (2, '8.2 Frontend Design Patterns'),
    (2, '8.3 Architectural Patterns'),
    (2, '8.4 Pattern Summary'),
    (1, '9. Unit Test Cases'),
    (2, '9.1 Backend Model Tests'),
    (2, '9.2 Backend Service Tests'),
    (2, '9.3 Backend Utility Tests'),
    (2, '9.4 Backend API Route Tests'),
    (2, '9.5 Frontend Component Tests'),
    (2, '9.6 Frontend API Service Tests'),
    (2, '9.7 Test Summary'),
    (1, '10. Architectural Design'),
    (2, '10.1 Architecture Overview'),
    (2, '10.2 Client Layer'),
    (2, '10.3 Server Layer'),
    (2, '10.4 Data Layer'),
    (2, '10.5 External Services'),
    (2, '10.6 File Storage'),
    (2, '10.7 Data Flow'),
    (2, '10.8 Deployment Architecture'),
    (2, '10.9 Architecture Diagram'),
    (2, '10.10 Key Architecture Decisions'),
    (1, '11. Conclusion'),
    (2, '11.1 Summary'),
    (2, '11.2 Project Scope Achievement'),
    (2, '11.3 Future Enhancements'),
    (2, '11.4 Technical Debt and Recommendations'),
    (2, '11.5 Final Remarks'),
]

# Add TOC entries with hyperlinks
for level, title in toc_entries:
    toc_para = doc.add_paragraph()
    bookmark_name = create_bookmark_id(title)

    # Set formatting based on level
    if level == 1:
        add_hyperlink(toc_para, title, bookmark_name, font_size=Pt(12), bold=True, indent=0)
    elif level == 2:
        add_hyperlink(toc_para, title, bookmark_name, font_size=Pt(11), bold=False, indent=0.25)
    else:  # level 3
        add_hyperlink(toc_para, title, bookmark_name, font_size=Pt(10), bold=False, indent=0.5)

# Add page break after TOC
doc.add_page_break()

# ============================================
# 1. INTRODUCTION
# ============================================
add_heading_with_bookmark(doc, '1. Introduction', level=1)

# 1.1 Purpose
add_heading_with_bookmark(doc, '1.1 Purpose', level=2)
doc.add_paragraph(
    'This Software Requirements Specification (SRS) document describes the functional and '
    'non-functional requirements for the Ask Holmes application. The purpose of this document '
    'is to provide a comprehensive overview of the system requirements, intended audience, '
    'and scope of the application. This document serves as a reference for developers, testers, '
    'and stakeholders involved in the development and deployment of the Ask Holmes application.'
)

# 1.2 Document Conventions
add_heading_with_bookmark(doc, '1.2 Document Conventions', level=2)
doc.add_paragraph(
    'This document follows standard SRS conventions. Requirements are identified using unique '
    'identifiers with the following prefixes:'
)
doc.add_paragraph('FR: Functional Requirements', style='List Bullet')
doc.add_paragraph('NFR: Non-Functional Requirements', style='List Bullet')
doc.add_paragraph('UX: User Experience Requirements', style='List Bullet')

# 1.3 Intended Audience
add_heading_with_bookmark(doc, '1.3 Intended Audience', level=2)
doc.add_paragraph('This document is intended for the following audiences:')
doc.add_paragraph('Developers: To understand the technical requirements and implement the system accordingly.', style='List Bullet')
doc.add_paragraph('Testers: To develop test cases and validate the system against the specified requirements.', style='List Bullet')
doc.add_paragraph('Project Managers: To track project progress and ensure requirements are being met.', style='List Bullet')
doc.add_paragraph('Stakeholders: To review and approve the system requirements before development begins.', style='List Bullet')

# 1.4 Product Scope
add_heading_with_bookmark(doc, '1.4 Product Scope', level=2)
doc.add_paragraph(
    'Ask Holmes is a web-based Retrieval-Augmented Generation (RAG) application that serves as '
    'an intelligent knowledge base for Sherlock Holmes literature. The system enables users to '
    'interact with classic Sherlock Holmes texts through a modern question-and-answer interface.'
)
doc.add_paragraph('The application provides the following core capabilities:')

doc.add_paragraph(
    'Question Management: Users can create, view, edit, and delete questions about Sherlock Holmes '
    'literature. Each question is securely associated with the authenticated user, ensuring data '
    'privacy and isolation.',
    style='List Bullet'
)
doc.add_paragraph(
    'AI-Powered Answers: The application leverages advanced AI technology to generate accurate '
    'answers by searching through indexed Sherlock Holmes books. Using semantic search and '
    'large language models, the system retrieves relevant passages and synthesizes comprehensive '
    'responses grounded in the source material.',
    style='List Bullet'
)
doc.add_paragraph(
    'Trusted Book Sources: Answers are derived exclusively from indexed Sherlock Holmes documents, '
    'ensuring authenticity and reliability. Users can trust that responses are based on the actual '
    'canon rather than fabricated information.',
    style='List Bullet'
)
doc.add_paragraph(
    'Manual Answer Entry: Users have the flexibility to provide their own answers to questions, '
    'edit AI-generated responses, or maintain a personal knowledge base of Sherlock Holmes facts.',
    style='List Bullet'
)

# 1.5 Product Overview
add_heading_with_bookmark(doc, '1.5 Product Overview', level=2)
doc.add_paragraph(
    'Ask Holmes addresses the need for an intelligent, user-friendly interface to explore and '
    'understand Sherlock Holmes literature. Traditional search methods often fail to provide '
    'contextual answers to complex questions about characters, plots, and themes. Ask Holmes '
    'solves this problem by combining modern AI techniques with classic literature.'
)

doc.add_paragraph('Key Features:')
doc.add_paragraph('Create questions through an intuitive modal-based interface', style='List Bullet')
doc.add_paragraph('View all personal questions in an organized list format', style='List Bullet')
doc.add_paragraph('Edit existing questions and answers at any time', style='List Bullet')
doc.add_paragraph('Delete questions that are no longer needed', style='List Bullet')
doc.add_paragraph('Request AI-generated answers using the "Ask Documents" feature', style='List Bullet')
doc.add_paragraph('Answers are sourced from trusted, indexed Sherlock Holmes books', style='List Bullet')
doc.add_paragraph('User data is isolated and secure through authentication', style='List Bullet')

# 1.6 References
add_heading_with_bookmark(doc, '1.6 References', level=2)
doc.add_paragraph('The following references were used in the development of this SRS:')
doc.add_paragraph('IEEE Std 830-1998: IEEE Recommended Practice for Software Requirements Specifications', style='List Bullet')
doc.add_paragraph('Project Scope Document: Ask Holmes Application Scope', style='List Bullet')
doc.add_paragraph('Technical Architecture Document: System Design and Components', style='List Bullet')

# 1.7 Definitions, Acronyms, and Abbreviations
add_heading_with_bookmark(doc, '1.7 Definitions, Acronyms, and Abbreviations', level=2)
doc.add_paragraph(
    'This section provides definitions for terms, acronyms, and abbreviations used throughout '
    'this document to ensure clarity and consistent understanding.'
)

# Acronyms subsection
add_heading_with_bookmark(doc, '1.7.1 Acronyms', level=3)

doc.add_paragraph('AI - Artificial Intelligence: The simulation of human intelligence processes by computer systems, including learning, reasoning, and self-correction.')
doc.add_paragraph('API - Application Programming Interface: A set of protocols and tools that allows different software applications to communicate with each other.')
doc.add_paragraph('ARIA - Accessible Rich Internet Applications: A set of attributes that define ways to make web content more accessible to people with disabilities.')
doc.add_paragraph('CORS - Cross-Origin Resource Sharing: A security mechanism that allows or restricts web applications running at one origin to access resources from a different origin.')
doc.add_paragraph('CPU - Central Processing Unit: The primary component of a computer that performs most of the processing inside the computer.')
doc.add_paragraph('CRUD - Create, Read, Update, Delete: The four basic operations of persistent storage in database applications.')
doc.add_paragraph('CSS - Cascading Style Sheets: A stylesheet language used to describe the presentation and visual formatting of HTML documents.')
doc.add_paragraph('DOM - Document Object Model: A programming interface for HTML documents that represents the page structure as a tree of objects.')
doc.add_paragraph('E2E - End-to-End: A testing methodology that tests the complete flow of an application from start to finish.')
doc.add_paragraph('GPU - Graphics Processing Unit: A specialized processor designed to accelerate graphics rendering and parallel processing tasks.')
doc.add_paragraph('HTML - HyperText Markup Language: The standard markup language for creating web pages and web applications.')
doc.add_paragraph('HTTP - HyperText Transfer Protocol: The foundation protocol used for transmitting data over the World Wide Web.')
doc.add_paragraph('HTTPS - HyperText Transfer Protocol Secure: An encrypted version of HTTP that provides secure communication over a computer network.')
doc.add_paragraph('IEEE - Institute of Electrical and Electronics Engineers: A professional association that develops standards for the electronics and computing industry.')
doc.add_paragraph('JSON - JavaScript Object Notation: A lightweight data interchange format that is easy for humans to read and write and easy for machines to parse.')
doc.add_paragraph('JSX - JavaScript XML: A syntax extension for JavaScript used in React that allows writing HTML-like code within JavaScript.')
doc.add_paragraph('JWT - JSON Web Token: A compact, URL-safe means of representing claims to be transferred between two parties for authentication.')
doc.add_paragraph('LLM - Large Language Model: An AI model trained on vast amounts of text data capable of understanding and generating human-like text.')
doc.add_paragraph('MFA - Multi-Factor Authentication: A security system that requires multiple methods of verification from independent categories of credentials.')
doc.add_paragraph('NLP - Natural Language Processing: A branch of AI that helps computers understand, interpret, and manipulate human language.')
doc.add_paragraph('NPM - Node Package Manager: A package manager for JavaScript that allows developers to share and reuse code packages.')
doc.add_paragraph('ORM - Object-Relational Mapping: A programming technique for converting data between incompatible type systems in object-oriented programming languages.')
doc.add_paragraph('PDF - Portable Document Format: A file format developed by Adobe to present documents independent of application software, hardware, and operating systems.')
doc.add_paragraph('PIP - Pip Installs Packages: The standard package manager for Python used to install and manage software packages.')
doc.add_paragraph('PWA - Progressive Web Application: A type of web application that uses modern web technologies to deliver app-like experiences.')
doc.add_paragraph('RAG - Retrieval-Augmented Generation: An AI framework that combines information retrieval with text generation to produce contextually relevant responses.')
doc.add_paragraph('RAM - Random Access Memory: A type of computer memory that can be read and changed in any order, used to store working data.')
doc.add_paragraph('REST - Representational State Transfer: An architectural style for designing networked applications using stateless client-server communication.')
doc.add_paragraph('SDK - Software Development Kit: A collection of software development tools in one installable package used to develop applications.')
doc.add_paragraph('SHA - Secure Hash Algorithm: A family of cryptographic hash functions used to verify data integrity.')
doc.add_paragraph('SPA - Single Page Application: A web application that dynamically rewrites the current page rather than loading entire new pages from the server.')
doc.add_paragraph('SQL - Structured Query Language: A domain-specific language used for managing and manipulating relational databases.')
doc.add_paragraph('SRS - Software Requirements Specification: A document that describes what the software will do and how it is expected to perform.')
doc.add_paragraph('SSL - Secure Sockets Layer: A cryptographic protocol designed to provide secure communication over a computer network.')
doc.add_paragraph('TLS - Transport Layer Security: A cryptographic protocol that provides secure communication over a network, successor to SSL.')
doc.add_paragraph('UI - User Interface: The visual elements through which users interact with a software application.')
doc.add_paragraph('URL - Uniform Resource Locator: The address used to access resources on the internet.')
doc.add_paragraph('UTF-8 - Unicode Transformation Format 8-bit: A variable-width character encoding capable of encoding all valid Unicode code points.')
doc.add_paragraph('UUID - Universally Unique Identifier: A 128-bit identifier used to uniquely identify information in computer systems.')
doc.add_paragraph('UX - User Experience: The overall experience of a person using a product, especially in terms of how easy or pleasing it is to use.')
doc.add_paragraph('WSGI - Web Server Gateway Interface: A specification for a universal interface between web servers and Python web applications.')
doc.add_paragraph('XSS - Cross-Site Scripting: A security vulnerability that allows attackers to inject malicious scripts into web pages viewed by other users.')

# Definitions subsection
add_heading_with_bookmark(doc, '1.7.2 Definitions', level=3)

doc.add_paragraph('Anthropic: An AI safety company that develops large language models, including the Claude family of AI assistants used in this application.')
doc.add_paragraph('Authentication: The process of verifying the identity of a user, typically through credentials such as email addresses or passwords.')
doc.add_paragraph('Authorization: The process of determining what permissions an authenticated user has to access specific resources or perform certain actions.')
doc.add_paragraph('Backend: The server-side portion of a web application that handles data processing, business logic, and database operations.')
doc.add_paragraph('Blueprint: In Flask, a way to organize a group of related views and other code, allowing modular application development.')
doc.add_paragraph('Chunk: A segment of text extracted from a document, typically 500 characters with overlap, used for efficient retrieval and processing.')
doc.add_paragraph('Claude: A family of large language models developed by Anthropic, with Claude Haiku being the fast and cost-effective variant used in this application.')
doc.add_paragraph('Component: In React, a reusable piece of UI that encapsulates its own structure, style, and behavior.')
doc.add_paragraph('Context API: A React feature that provides a way to pass data through the component tree without having to pass props manually at every level.')
doc.add_paragraph('Cosine Similarity: A mathematical measure of similarity between two vectors, used to find semantically similar text chunks.')
doc.add_paragraph('Docker: A platform for developing, shipping, and running applications in isolated containers.')
doc.add_paragraph('Embedding: A numerical representation of text as a vector in high-dimensional space, capturing semantic meaning for similarity comparisons.')
doc.add_paragraph('Endpoint: A specific URL path in an API that accepts requests and returns responses for a particular resource or operation.')
doc.add_paragraph('Flask: A lightweight Python web framework used for building web applications and RESTful APIs.')
doc.add_paragraph('Frontend: The client-side portion of a web application that users interact with directly, typically running in a web browser.')
doc.add_paragraph('Gunicorn: A Python WSGI HTTP server commonly used to run Python web applications in production environments.')
doc.add_paragraph('Hook: In React, a function that lets you use state and other React features in functional components.')
doc.add_paragraph('Indexing: The process of extracting text from documents, generating embeddings, and storing them in a database for efficient retrieval.')
doc.add_paragraph('IVFFlat: An index type in pgvector that uses inverted file indexing for approximate nearest neighbor search.')
doc.add_paragraph('Modal: A dialog box or popup window that appears on top of the main content, requiring user interaction before returning to the main view.')
doc.add_paragraph('Nginx: A high-performance web server and reverse proxy server commonly used to serve static files and proxy requests to backend applications.')
doc.add_paragraph('Node.js: A JavaScript runtime built on Chrome\'s V8 JavaScript engine, used for building server-side applications.')
doc.add_paragraph('pgvector: A PostgreSQL extension that adds support for vector similarity search, enabling semantic search capabilities.')
doc.add_paragraph('PostgreSQL: An open-source relational database management system known for reliability, feature robustness, and extensibility.')
doc.add_paragraph('Props: In React, read-only properties passed from parent components to child components to configure their behavior and display.')
doc.add_paragraph('PyPDF2: A Python library used for extracting text, merging, splitting, and manipulating PDF files.')
doc.add_paragraph('Query: A request for data or information from a database, or in RAG context, a user\'s question submitted for answer generation.')
doc.add_paragraph('React: A JavaScript library for building user interfaces, particularly single-page applications, developed by Facebook.')
doc.add_paragraph('Semantic Search: A search technique that understands the contextual meaning of terms to find relevant results beyond exact keyword matching.')
doc.add_paragraph('Sentence Transformers: A Python library for computing dense vector representations of sentences and paragraphs.')
doc.add_paragraph('SQLAlchemy: A Python SQL toolkit and Object-Relational Mapping library for database operations.')
doc.add_paragraph('State: In React, an object that holds data that may change over the lifetime of a component, triggering re-renders when updated.')
doc.add_paragraph('Toast: A brief notification message that appears temporarily on the screen to provide feedback about an operation.')
doc.add_paragraph('Token: In LLM context, a unit of text (word, subword, or character) that the model processes; used for billing and context limits.')
doc.add_paragraph('Vector: A mathematical representation of data as a list of numbers, used in this application for semantic similarity comparisons.')
doc.add_paragraph('Vector Database: A database optimized for storing and querying high-dimensional vector embeddings for similarity search.')
doc.add_paragraph('Viewport: The visible area of a web page in the browser window, which varies based on the device and screen size.')

# Technical Terms subsection
add_heading_with_bookmark(doc, '1.7.3 Technical Terms', level=3)

doc.add_paragraph('all-MiniLM-L6-v2: A sentence transformer model that produces 384-dimensional embeddings, balancing performance and efficiency.')
doc.add_paragraph('Connection Pooling: A technique of maintaining a cache of database connections that can be reused for future requests.')
doc.add_paragraph('CRUD Operations: The four basic database operations: Create (insert new records), Read (retrieve records), Update (modify records), Delete (remove records).')
doc.add_paragraph('Data Isolation: A security principle ensuring that each user can only access their own data and cannot view or modify other users\' data.')
doc.add_paragraph('Dependency Injection: A design pattern where dependencies are provided to a component rather than created by the component itself.')
doc.add_paragraph('Environment Variable: A dynamic value stored outside the application code that can affect the behavior of running processes.')
doc.add_paragraph('Factory Pattern: A design pattern that provides an interface for creating objects without specifying their concrete classes.')
doc.add_paragraph('Foreign Key: A field in a database table that references the primary key of another table, establishing a relationship between tables.')
doc.add_paragraph('Horizontal Scaling: Adding more machines or instances to distribute load, as opposed to vertical scaling which adds more power to existing machines.')
doc.add_paragraph('Idempotent: An operation that produces the same result regardless of how many times it is executed.')
doc.add_paragraph('Middleware: Software that acts as a bridge between an operating system or database and applications, handling requests and responses.')
doc.add_paragraph('Primary Key: A unique identifier for each record in a database table, ensuring no duplicate entries.')
doc.add_paragraph('Rate Limiting: A technique used to control the number of requests a client can make to an API within a specified time period.')
doc.add_paragraph('Reverse Proxy: A server that sits in front of web servers and forwards client requests to those servers.')
doc.add_paragraph('Service Layer: An architectural pattern that defines an application\'s boundary with a layer of services that establishes business logic.')
doc.add_paragraph('Singleton Pattern: A design pattern that restricts the instantiation of a class to a single instance.')
doc.add_paragraph('Stateless: A system design where no client context is stored on the server between requests; each request contains all information needed to process it.')
doc.add_paragraph('Three-Tier Architecture: A software architecture pattern that separates an application into three logical layers: presentation, application, and data.')
doc.add_paragraph('Timeout: A specified period after which an operation is automatically terminated if not completed.')
doc.add_paragraph('Timestamp: A sequence of characters representing the date and time at which a certain event occurred.')

# ============================================
# 2. PROJECT REQUIREMENTS
# ============================================
add_heading_with_bookmark(doc, '2. Project Requirements', level=1)
doc.add_paragraph(
    'This chapter defines the detailed functional and non-functional requirements for the '
    'Ask Holmes application. Each requirement is assigned a unique identifier for traceability '
    'and testing purposes.'
)

# ============================================
# 2.1 FUNCTIONAL REQUIREMENTS
# ============================================
add_heading_with_bookmark(doc, '2.1 Functional Requirements', level=2)
doc.add_paragraph(
    'Functional requirements describe the specific behaviors, features, and functions that the '
    'Ask Holmes application must provide. These requirements are organized by functional area.'
)

# 2.1.1 User Authentication Requirements
add_heading_with_bookmark(doc, '2.1.1 User Authentication Requirements', level=3)

doc.add_paragraph('FR-AUTH-001: User Registration')
doc.add_paragraph('Description: The system shall allow new users to register by providing a valid email address.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User enters email address in the registration form', style='List Bullet')
doc.add_paragraph('2. System validates email format (contains @ and valid domain)', style='List Bullet')
doc.add_paragraph('3. System checks if email is not already registered', style='List Bullet')
doc.add_paragraph('4. System creates new user account with unique UUID', style='List Bullet')
doc.add_paragraph('5. System displays success message upon successful registration', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-AUTH-002: User Login')
doc.add_paragraph('Description: The system shall allow registered users to log in using their email address.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User enters registered email address', style='List Bullet')
doc.add_paragraph('2. System validates email exists in database', style='List Bullet')
doc.add_paragraph('3. System authenticates user and creates session', style='List Bullet')
doc.add_paragraph('4. System stores user ID in browser local storage', style='List Bullet')
doc.add_paragraph('5. System redirects user to home page', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-AUTH-003: User Logout')
doc.add_paragraph('Description: The system shall allow authenticated users to log out of the application.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User clicks logout button in header', style='List Bullet')
doc.add_paragraph('2. System clears user session from local storage', style='List Bullet')
doc.add_paragraph('3. System redirects user to login page', style='List Bullet')
doc.add_paragraph('4. System prevents access to protected pages without re-authentication', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-AUTH-004: Session Persistence')
doc.add_paragraph('Description: The system shall maintain user session across browser refreshes and new tabs.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User ID is stored in browser local storage upon login', style='List Bullet')
doc.add_paragraph('2. Application checks local storage on page load', style='List Bullet')
doc.add_paragraph('3. If valid user ID exists, user remains authenticated', style='List Bullet')
doc.add_paragraph('4. User can open multiple tabs without re-authenticating', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

# 2.1.2 Question Management Requirements
add_heading_with_bookmark(doc, '2.1.2 Question Management Requirements', level=3)

doc.add_paragraph('FR-Q-001: Create Question')
doc.add_paragraph('Description: The system shall allow authenticated users to create new questions through a modal interface.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User clicks "Add Question" button on home page', style='List Bullet')
doc.add_paragraph('2. Modal dialog opens with question input form', style='List Bullet')
doc.add_paragraph('3. User enters question text (required field)', style='List Bullet')
doc.add_paragraph('4. User optionally enters answer text', style='List Bullet')
doc.add_paragraph('5. User clicks "Save" button to submit', style='List Bullet')
doc.add_paragraph('6. System validates question is not empty', style='List Bullet')
doc.add_paragraph('7. System saves question with user ID and timestamp', style='List Bullet')
doc.add_paragraph('8. Modal closes and question appears in list', style='List Bullet')
doc.add_paragraph('9. Success toast notification is displayed', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-Q-002: View Questions List')
doc.add_paragraph('Description: The system shall display all questions belonging to the authenticated user.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Home page displays list of user\'s questions', style='List Bullet')
doc.add_paragraph('2. Questions are sorted by creation date (newest first)', style='List Bullet')
doc.add_paragraph('3. Each question shows question text and answer preview', style='List Bullet')
doc.add_paragraph('4. Questions without answers show "No answer yet" indicator', style='List Bullet')
doc.add_paragraph('5. Only questions belonging to current user are displayed', style='List Bullet')
doc.add_paragraph('6. Empty state message shown when no questions exist', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-Q-003: Edit Question')
doc.add_paragraph('Description: The system shall allow users to edit their existing questions.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User clicks edit option from question menu', style='List Bullet')
doc.add_paragraph('2. Modal opens pre-populated with existing question and answer', style='List Bullet')
doc.add_paragraph('3. User modifies question text and/or answer', style='List Bullet')
doc.add_paragraph('4. User clicks "Save" to submit changes', style='List Bullet')
doc.add_paragraph('5. System validates question is not empty', style='List Bullet')
doc.add_paragraph('6. System updates question with new updated_at timestamp', style='List Bullet')
doc.add_paragraph('7. Modal closes and list reflects changes', style='List Bullet')
doc.add_paragraph('8. Success toast notification is displayed', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-Q-004: Delete Question')
doc.add_paragraph('Description: The system shall allow users to permanently delete their questions.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User clicks delete option from question menu', style='List Bullet')
doc.add_paragraph('2. Confirmation dialog appears with warning message', style='List Bullet')
doc.add_paragraph('3. User confirms deletion by clicking "Delete" button', style='List Bullet')
doc.add_paragraph('4. System permanently removes question and associated answer', style='List Bullet')
doc.add_paragraph('5. Question is removed from the displayed list', style='List Bullet')
doc.add_paragraph('6. Success toast notification is displayed', style='List Bullet')
doc.add_paragraph('7. User can cancel deletion by clicking "Cancel" or pressing Escape', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-Q-005: Question Validation')
doc.add_paragraph('Description: The system shall validate question input before saving.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Question text is required and cannot be empty', style='List Bullet')
doc.add_paragraph('2. Question must contain at least one non-whitespace character', style='List Bullet')
doc.add_paragraph('3. Leading and trailing whitespace is trimmed', style='List Bullet')
doc.add_paragraph('4. Answer field is optional (can be empty)', style='List Bullet')
doc.add_paragraph('5. Validation error message displayed for invalid input', style='List Bullet')
doc.add_paragraph('6. Save button disabled when form is invalid', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

# 2.1.3 Answer Management Requirements
add_heading_with_bookmark(doc, '2.1.3 Answer Management Requirements', level=3)

doc.add_paragraph('FR-A-001: Manual Answer Entry')
doc.add_paragraph('Description: The system shall allow users to manually enter answers to their questions.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Answer textarea is available in question creation/edit modal', style='List Bullet')
doc.add_paragraph('2. User can enter free-form text as answer', style='List Bullet')
doc.add_paragraph('3. Answer is saved with the question', style='List Bullet')
doc.add_paragraph('4. Empty answer is stored as null in database', style='List Bullet')
doc.add_paragraph('5. Answer can be added to questions that previously had none', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-A-002: AI Answer Generation')
doc.add_paragraph('Description: The system shall generate AI-powered answers using the RAG pipeline.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. "Ask Documents" button is visible in question form', style='List Bullet')
doc.add_paragraph('2. Button is disabled when question field is empty', style='List Bullet')
doc.add_paragraph('3. Clicking button triggers RAG query with question text', style='List Bullet')
doc.add_paragraph('4. Loading indicator shown during answer generation', style='List Bullet')
doc.add_paragraph('5. Generated answer is populated in answer textarea', style='List Bullet')
doc.add_paragraph('6. User can edit generated answer before saving', style='List Bullet')
doc.add_paragraph('7. Error message displayed if generation fails', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-A-003: Answer Editing')
doc.add_paragraph('Description: The system shall allow users to edit both manual and AI-generated answers.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Existing answers are editable in the edit modal', style='List Bullet')
doc.add_paragraph('2. AI-generated answers can be modified before saving', style='List Bullet')
doc.add_paragraph('3. Answers can be cleared (set to empty)', style='List Bullet')
doc.add_paragraph('4. Changes are saved when user clicks "Save"', style='List Bullet')
doc.add_paragraph('5. Updated timestamp reflects modification time', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('FR-A-004: Answer Regeneration')
doc.add_paragraph('Description: The system shall allow users to regenerate AI answers.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. "Ask Documents" button can be clicked multiple times', style='List Bullet')
doc.add_paragraph('2. Each click generates a new answer based on current question', style='List Bullet')
doc.add_paragraph('3. New answer replaces previous content in textarea', style='List Bullet')
doc.add_paragraph('4. Saved answers are not affected until user clicks "Save"', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# 2.1.4 RAG Pipeline Requirements
add_heading_with_bookmark(doc, '2.1.4 RAG Pipeline Requirements', level=3)

doc.add_paragraph('FR-RAG-001: Document Upload')
doc.add_paragraph('Description: The system shall allow uploading PDF documents for indexing.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. File upload button accepts PDF files only', style='List Bullet')
doc.add_paragraph('2. Selected file is uploaded to server books folder', style='List Bullet')
doc.add_paragraph('3. Upload progress indicator is displayed', style='List Bullet')
doc.add_paragraph('4. Success message shown upon completion', style='List Bullet')
doc.add_paragraph('5. Error message shown for invalid file types', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('FR-RAG-002: Document Indexing')
doc.add_paragraph('Description: The system shall process and index uploaded PDF documents.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. System extracts text content from PDF files', style='List Bullet')
doc.add_paragraph('2. Text is split into chunks of 500 characters with 50-character overlap', style='List Bullet')
doc.add_paragraph('3. Each chunk is converted to 384-dimensional vector embedding', style='List Bullet')
doc.add_paragraph('4. Chunks and embeddings are stored in database', style='List Bullet')
doc.add_paragraph('5. Document metadata (filename, hash, chunk count) is recorded', style='List Bullet')
doc.add_paragraph('6. Duplicate documents are detected via file hash', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-RAG-003: Semantic Search')
doc.add_paragraph('Description: The system shall perform semantic similarity search on indexed documents.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. User question is converted to vector embedding', style='List Bullet')
doc.add_paragraph('2. System searches for most similar document chunks', style='List Bullet')
doc.add_paragraph('3. Top 3 most relevant chunks are retrieved', style='List Bullet')
doc.add_paragraph('4. Cosine similarity is used for matching', style='List Bullet')
doc.add_paragraph('5. Search completes within acceptable time limit', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-RAG-004: Answer Generation')
doc.add_paragraph('Description: The system shall generate answers using retrieved context and LLM.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Retrieved chunks are assembled as context', style='List Bullet')
doc.add_paragraph('2. Question and context are sent to Claude Haiku LLM', style='List Bullet')
doc.add_paragraph('3. LLM generates natural language answer', style='List Bullet')
doc.add_paragraph('4. Answer is grounded in the provided context', style='List Bullet')
doc.add_paragraph('5. Response is returned within 60-second timeout', style='List Bullet')
doc.add_paragraph('6. Error handling for API failures', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-RAG-005: Document Management')
doc.add_paragraph('Description: The system shall allow viewing and deleting indexed documents.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. List of indexed documents is viewable', style='List Bullet')
doc.add_paragraph('2. Document details show filename and chunk count', style='List Bullet')
doc.add_paragraph('3. Documents can be deleted from the index', style='List Bullet')
doc.add_paragraph('4. Deleting document removes all associated chunks', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# 2.1.5 User Interface Requirements
add_heading_with_bookmark(doc, '2.1.5 User Interface Requirements', level=3)

doc.add_paragraph('FR-UI-001: Responsive Layout')
doc.add_paragraph('Description: The system shall provide a responsive layout for desktop and tablet devices.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Layout adapts to screen sizes 1024px and above', style='List Bullet')
doc.add_paragraph('2. All features are accessible on desktop browsers', style='List Bullet')
doc.add_paragraph('3. Content is readable without horizontal scrolling', style='List Bullet')
doc.add_paragraph('4. Interactive elements are appropriately sized', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('FR-UI-002: Modal Dialogs')
doc.add_paragraph('Description: The system shall use modal dialogs for question creation and editing.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Modal opens centered on screen with backdrop overlay', style='List Bullet')
doc.add_paragraph('2. Modal can be closed by clicking X button', style='List Bullet')
doc.add_paragraph('3. Modal can be closed by pressing Escape key', style='List Bullet')
doc.add_paragraph('4. Modal can be closed by clicking backdrop', style='List Bullet')
doc.add_paragraph('5. Focus is trapped within modal when open', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('FR-UI-003: Toast Notifications')
doc.add_paragraph('Description: The system shall display toast notifications for user feedback.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Success messages shown for successful operations', style='List Bullet')
doc.add_paragraph('2. Error messages shown for failed operations', style='List Bullet')
doc.add_paragraph('3. Toasts appear briefly and auto-dismiss', style='List Bullet')
doc.add_paragraph('4. Toasts are positioned consistently on screen', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

doc.add_paragraph('FR-UI-004: Loading States')
doc.add_paragraph('Description: The system shall display loading indicators during async operations.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Loading spinner shown while fetching questions', style='List Bullet')
doc.add_paragraph('2. Loading indicator shown during AI answer generation', style='List Bullet')
doc.add_paragraph('3. Buttons are disabled during loading states', style='List Bullet')
doc.add_paragraph('4. Loading states prevent duplicate submissions', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('FR-UI-005: Error Handling')
doc.add_paragraph('Description: The system shall gracefully handle and display errors.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Network errors display appropriate message', style='List Bullet')
doc.add_paragraph('2. Validation errors show inline feedback', style='List Bullet')
doc.add_paragraph('3. API errors are caught and displayed to user', style='List Bullet')
doc.add_paragraph('4. Application does not crash on errors', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

# 2.1.6 Data Management Requirements
add_heading_with_bookmark(doc, '2.1.6 Data Management Requirements', level=3)

doc.add_paragraph('FR-DATA-001: Data Isolation')
doc.add_paragraph('Description: The system shall ensure users can only access their own data.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Questions are associated with user ID on creation', style='List Bullet')
doc.add_paragraph('2. API endpoints filter data by authenticated user', style='List Bullet')
doc.add_paragraph('3. Attempts to access other users\' data return 403 error', style='List Bullet')
doc.add_paragraph('4. User ID is validated on all protected endpoints', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-DATA-002: Data Persistence')
doc.add_paragraph('Description: The system shall persist all user data to the database.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Questions are stored in PostgreSQL database', style='List Bullet')
doc.add_paragraph('2. Data survives application restarts', style='List Bullet')
doc.add_paragraph('3. Timestamps are recorded for all records', style='List Bullet')
doc.add_paragraph('4. UUIDs are used for primary keys', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('FR-DATA-003: Cascade Deletion')
doc.add_paragraph('Description: The system shall properly handle related data when deleting records.', style='List Bullet')
doc.add_paragraph('Acceptance Criteria:', style='List Bullet')
doc.add_paragraph('1. Deleting a user removes all associated questions', style='List Bullet')
doc.add_paragraph('2. Deleting a document removes all associated chunks', style='List Bullet')
doc.add_paragraph('3. No orphaned records remain after deletion', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

# ============================================
# 2.2 NON-FUNCTIONAL REQUIREMENTS
# ============================================
add_heading_with_bookmark(doc, '2.2 Non-Functional Requirements', level=2)
doc.add_paragraph(
    'Non-functional requirements define the quality attributes, constraints, and characteristics '
    'that the Ask Holmes application must exhibit. These requirements address performance, '
    'security, usability, and other system qualities.'
)

# 2.2.1 Performance Requirements
add_heading_with_bookmark(doc, '2.2.1 Performance Requirements', level=3)

doc.add_paragraph('NFR-PERF-001: API Response Time')
doc.add_paragraph('Description: The system shall respond to API requests within acceptable time limits.', style='List Bullet')
doc.add_paragraph('Requirement: CRUD operations shall complete within 2 seconds under normal load.', style='List Bullet')
doc.add_paragraph('Measurement: Server response time measured from request receipt to response sent.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-PERF-002: RAG Query Response Time')
doc.add_paragraph('Description: The system shall generate AI answers within acceptable time limits.', style='List Bullet')
doc.add_paragraph('Requirement: RAG answer generation shall complete within 60 seconds.', style='List Bullet')
doc.add_paragraph('Measurement: Time from query submission to answer receipt.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-PERF-003: Page Load Time')
doc.add_paragraph('Description: The system shall load pages within acceptable time limits.', style='List Bullet')
doc.add_paragraph('Requirement: Initial page load shall complete within 3 seconds on broadband connection.', style='List Bullet')
doc.add_paragraph('Measurement: Time from navigation to fully rendered page.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-PERF-004: Concurrent Users')
doc.add_paragraph('Description: The system shall support multiple simultaneous users.', style='List Bullet')
doc.add_paragraph('Requirement: System shall support 100 concurrent users under normal operation.', style='List Bullet')
doc.add_paragraph('Maximum Capacity: System shall handle up to 500 concurrent users at peak load.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-PERF-005: Database Query Performance')
doc.add_paragraph('Description: The system shall execute database queries efficiently.', style='List Bullet')
doc.add_paragraph('Requirement: Question list queries shall complete within 500ms.', style='List Bullet')
doc.add_paragraph('Requirement: Vector similarity search shall complete within 2 seconds.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

# 2.2.2 Security Requirements
add_heading_with_bookmark(doc, '2.2.2 Security Requirements', level=3)

doc.add_paragraph('NFR-SEC-001: Data Isolation')
doc.add_paragraph('Description: The system shall enforce strict data isolation between users.', style='List Bullet')
doc.add_paragraph('Requirement: Users shall only access their own questions and answers.', style='List Bullet')
doc.add_paragraph('Implementation: Server-side validation of user ownership on all requests.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-SEC-002: Input Validation')
doc.add_paragraph('Description: The system shall validate and sanitize all user inputs.', style='List Bullet')
doc.add_paragraph('Requirement: All inputs shall be validated for type, length, and format.', style='List Bullet')
doc.add_paragraph('Requirement: Input sanitization shall prevent XSS attacks.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-SEC-003: SQL Injection Prevention')
doc.add_paragraph('Description: The system shall prevent SQL injection attacks.', style='List Bullet')
doc.add_paragraph('Requirement: All database queries shall use parameterized statements via ORM.', style='List Bullet')
doc.add_paragraph('Implementation: SQLAlchemy ORM with bound parameters.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-SEC-004: Secure Communication')
doc.add_paragraph('Description: The system shall encrypt all network communications.', style='List Bullet')
doc.add_paragraph('Requirement: All client-server communication shall use HTTPS/TLS 1.2+.', style='List Bullet')
doc.add_paragraph('Requirement: API keys shall never be exposed to client-side code.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-SEC-005: API Key Protection')
doc.add_paragraph('Description: The system shall securely manage API credentials.', style='List Bullet')
doc.add_paragraph('Requirement: Anthropic API keys shall be stored as environment variables.', style='List Bullet')
doc.add_paragraph('Requirement: API keys shall not be committed to version control.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-SEC-006: Error Message Security')
doc.add_paragraph('Description: The system shall not expose sensitive information in error messages.', style='List Bullet')
doc.add_paragraph('Requirement: Error messages shall not reveal system internals or stack traces to users.', style='List Bullet')
doc.add_paragraph('Requirement: Detailed errors shall only be logged server-side.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

# 2.2.3 Reliability Requirements
add_heading_with_bookmark(doc, '2.2.3 Reliability Requirements', level=3)

doc.add_paragraph('NFR-REL-001: Data Persistence')
doc.add_paragraph('Description: The system shall reliably persist all user data.', style='List Bullet')
doc.add_paragraph('Requirement: Data shall survive application restarts and server reboots.', style='List Bullet')
doc.add_paragraph('Implementation: PostgreSQL database with ACID compliance.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-REL-002: Error Recovery')
doc.add_paragraph('Description: The system shall gracefully handle errors and recover when possible.', style='List Bullet')
doc.add_paragraph('Requirement: Application shall not crash due to user input errors.', style='List Bullet')
doc.add_paragraph('Requirement: Network failures shall display appropriate messages without data loss.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-REL-003: External Service Failure')
doc.add_paragraph('Description: The system shall handle external service failures gracefully.', style='List Bullet')
doc.add_paragraph('Requirement: Anthropic API failures shall display user-friendly error messages.', style='List Bullet')
doc.add_paragraph('Requirement: Core functionality (CRUD) shall work without external services.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-REL-004: Data Integrity')
doc.add_paragraph('Description: The system shall maintain data integrity at all times.', style='List Bullet')
doc.add_paragraph('Requirement: Database operations shall be atomic (all-or-nothing).', style='List Bullet')
doc.add_paragraph('Requirement: Foreign key constraints shall prevent orphaned records.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

# 2.2.4 Usability Requirements
add_heading_with_bookmark(doc, '2.2.4 Usability Requirements', level=3)

doc.add_paragraph('NFR-USA-001: Intuitive Interface')
doc.add_paragraph('Description: The system shall provide an intuitive and easy-to-use interface.', style='List Bullet')
doc.add_paragraph('Requirement: New users shall be able to create a question within 2 minutes without training.', style='List Bullet')
doc.add_paragraph('Requirement: Primary actions shall be discoverable without documentation.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-USA-002: Consistent Design')
doc.add_paragraph('Description: The system shall maintain consistent design patterns throughout.', style='List Bullet')
doc.add_paragraph('Requirement: UI elements shall follow consistent styling and behavior.', style='List Bullet')
doc.add_paragraph('Requirement: Navigation patterns shall be predictable across pages.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-USA-003: Feedback')
doc.add_paragraph('Description: The system shall provide clear feedback for all user actions.', style='List Bullet')
doc.add_paragraph('Requirement: Success and error states shall be clearly communicated.', style='List Bullet')
doc.add_paragraph('Requirement: Loading states shall indicate ongoing operations.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-USA-004: Error Messages')
doc.add_paragraph('Description: The system shall display helpful error messages.', style='List Bullet')
doc.add_paragraph('Requirement: Error messages shall explain what went wrong.', style='List Bullet')
doc.add_paragraph('Requirement: Error messages shall suggest corrective actions when possible.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# 2.2.5 Accessibility Requirements
add_heading_with_bookmark(doc, '2.2.5 Accessibility Requirements', level=3)

doc.add_paragraph('NFR-ACC-001: Keyboard Navigation')
doc.add_paragraph('Description: The system shall support keyboard navigation.', style='List Bullet')
doc.add_paragraph('Requirement: All interactive elements shall be accessible via keyboard.', style='List Bullet')
doc.add_paragraph('Requirement: Tab order shall follow logical reading order.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-ACC-002: Screen Reader Support')
doc.add_paragraph('Description: The system shall be compatible with screen readers.', style='List Bullet')
doc.add_paragraph('Requirement: Interactive elements shall have appropriate ARIA labels.', style='List Bullet')
doc.add_paragraph('Requirement: Dynamic content changes shall be announced.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

doc.add_paragraph('NFR-ACC-003: Color Contrast')
doc.add_paragraph('Description: The system shall maintain sufficient color contrast.', style='List Bullet')
doc.add_paragraph('Requirement: Text shall have minimum 4.5:1 contrast ratio against background.', style='List Bullet')
doc.add_paragraph('Requirement: Interactive elements shall be distinguishable without color alone.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

doc.add_paragraph('NFR-ACC-004: Focus Indicators')
doc.add_paragraph('Description: The system shall provide visible focus indicators.', style='List Bullet')
doc.add_paragraph('Requirement: Focused elements shall have visible outline or highlight.', style='List Bullet')
doc.add_paragraph('Requirement: Focus state shall not rely on color alone.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# 2.2.6 Maintainability Requirements
add_heading_with_bookmark(doc, '2.2.6 Maintainability Requirements', level=3)

doc.add_paragraph('NFR-MAINT-001: Code Organization')
doc.add_paragraph('Description: The system shall follow established code organization patterns.', style='List Bullet')
doc.add_paragraph('Requirement: Backend shall use service layer pattern for business logic.', style='List Bullet')
doc.add_paragraph('Requirement: Frontend shall use component-based architecture.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-MAINT-002: Separation of Concerns')
doc.add_paragraph('Description: The system shall maintain clear separation of concerns.', style='List Bullet')
doc.add_paragraph('Requirement: Presentation, business logic, and data access shall be separated.', style='List Bullet')
doc.add_paragraph('Requirement: API routes shall delegate to service layer.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-MAINT-003: Configuration Management')
doc.add_paragraph('Description: The system shall externalize configuration.', style='List Bullet')
doc.add_paragraph('Requirement: Environment-specific values shall be configurable via environment variables.', style='List Bullet')
doc.add_paragraph('Requirement: No hardcoded credentials or environment-specific values in code.', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

# 2.2.7 Scalability Requirements
add_heading_with_bookmark(doc, '2.2.7 Scalability Requirements', level=3)

doc.add_paragraph('NFR-SCALE-001: Horizontal Scalability')
doc.add_paragraph('Description: The system shall support horizontal scaling.', style='List Bullet')
doc.add_paragraph('Requirement: Application shall be stateless to allow multiple instances.', style='List Bullet')
doc.add_paragraph('Requirement: Session data shall be stored externally (database/local storage).', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

doc.add_paragraph('NFR-SCALE-002: Database Scalability')
doc.add_paragraph('Description: The system shall use scalable database patterns.', style='List Bullet')
doc.add_paragraph('Requirement: Database queries shall use proper indexing.', style='List Bullet')
doc.add_paragraph('Requirement: Connection pooling shall be implemented.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-SCALE-003: Document Volume')
doc.add_paragraph('Description: The system shall handle growing document volumes.', style='List Bullet')
doc.add_paragraph('Requirement: Vector search shall remain performant with up to 100,000 chunks.', style='List Bullet')
doc.add_paragraph('Requirement: IVFFlat index shall be used for efficient similarity search.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# 2.2.8 Compatibility Requirements
add_heading_with_bookmark(doc, '2.2.8 Compatibility Requirements', level=3)

doc.add_paragraph('NFR-COMPAT-001: Browser Compatibility')
doc.add_paragraph('Description: The system shall support modern web browsers.', style='List Bullet')
doc.add_paragraph('Supported Browsers:', style='List Bullet')
doc.add_paragraph('- Google Chrome 100+', style='List Bullet')
doc.add_paragraph('- Mozilla Firefox 100+', style='List Bullet')
doc.add_paragraph('- Microsoft Edge 100+', style='List Bullet')
doc.add_paragraph('- Apple Safari 15+', style='List Bullet')
doc.add_paragraph('Priority: High', style='List Bullet')

doc.add_paragraph('NFR-COMPAT-002: Screen Resolution')
doc.add_paragraph('Description: The system shall support standard screen resolutions.', style='List Bullet')
doc.add_paragraph('Requirement: Minimum supported resolution shall be 1024x768 pixels.', style='List Bullet')
doc.add_paragraph('Requirement: Layout shall be optimized for desktop and tablet viewports.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-COMPAT-003: JavaScript Requirement')
doc.add_paragraph('Description: The system shall require JavaScript for full functionality.', style='List Bullet')
doc.add_paragraph('Requirement: Application requires JavaScript to be enabled in browser.', style='List Bullet')
doc.add_paragraph('Requirement: Graceful message displayed if JavaScript is disabled.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# 2.2.9 Deployment Requirements
add_heading_with_bookmark(doc, '2.2.9 Deployment Requirements', level=3)

doc.add_paragraph('NFR-DEPLOY-001: Containerization')
doc.add_paragraph('Description: The system shall support containerized deployment.', style='List Bullet')
doc.add_paragraph('Requirement: Application shall be deployable via Docker containers.', style='List Bullet')
doc.add_paragraph('Requirement: Docker Compose shall orchestrate multi-container setup.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-DEPLOY-002: Environment Configuration')
doc.add_paragraph('Description: The system shall support multiple deployment environments.', style='List Bullet')
doc.add_paragraph('Requirement: Configuration shall be managed via environment variables.', style='List Bullet')
doc.add_paragraph('Requirement: Same container image shall work in different environments.', style='List Bullet')
doc.add_paragraph('Priority: Medium', style='List Bullet')

doc.add_paragraph('NFR-DEPLOY-003: Database Migration')
doc.add_paragraph('Description: The system shall support database schema management.', style='List Bullet')
doc.add_paragraph('Requirement: Database schema shall be initialized via SQL scripts.', style='List Bullet')
doc.add_paragraph('Requirement: Schema changes shall be versioned and repeatable.', style='List Bullet')
doc.add_paragraph('Priority: Low', style='List Bullet')

# ============================================
# 3. USE CASE DIAGRAM
# ============================================
add_heading_with_bookmark(doc, '3. Use Case Diagram', level=1)
doc.add_paragraph(
    'This chapter presents the Use Case Diagram for the Ask Holmes application, illustrating '
    'the interactions between actors and the system. The diagram identifies all primary and '
    'secondary actors, use cases organized by functional packages, and the relationships '
    'between them.'
)

# 3.1 Actors
add_heading_with_bookmark(doc, '3.1 Actors', level=2)
doc.add_paragraph(
    'The following actors interact with the Ask Holmes application:'
)

doc.add_paragraph('Actor 1: User (Primary Actor)')
doc.add_paragraph('Type: Human', style='List Bullet')
doc.add_paragraph('Description: A registered user of the Ask Holmes application who can create, view, edit, and delete questions, as well as request AI-generated answers from indexed Sherlock Holmes literature.', style='List Bullet')
doc.add_paragraph('Goals: Manage personal Q&A knowledge base, obtain accurate answers about Sherlock Holmes literature.', style='List Bullet')
doc.add_paragraph('Characteristics: Low to moderate technical expertise, familiar with web applications.', style='List Bullet')

doc.add_paragraph('Actor 2: Administrator (Primary Actor)')
doc.add_paragraph('Type: Human', style='List Bullet')
doc.add_paragraph('Description: A system administrator who has all User capabilities plus the ability to upload, index, and manage PDF documents that serve as the knowledge base for the RAG system.', style='List Bullet')
doc.add_paragraph('Goals: Maintain the document repository, ensure indexed content is accurate and up-to-date.', style='List Bullet')
doc.add_paragraph('Characteristics: Moderate to high technical expertise, understands document management.', style='List Bullet')
doc.add_paragraph('Relationship: Inherits from User (specialization)', style='List Bullet')

doc.add_paragraph('Actor 3: Claude API (Secondary Actor)')
doc.add_paragraph('Type: External System', style='List Bullet')
doc.add_paragraph('Description: Anthropic\'s Claude Haiku large language model API that generates natural language answers based on retrieved document context.', style='List Bullet')
doc.add_paragraph('Goals: Provide accurate, contextual answers grounded in the source material.', style='List Bullet')
doc.add_paragraph('Interface: HTTPS REST API', style='List Bullet')

doc.add_paragraph('Actor 4: Embedding Model (Secondary Actor)')
doc.add_paragraph('Type: External System', style='List Bullet')
doc.add_paragraph('Description: Sentence Transformers all-MiniLM-L6-v2 model that generates 384-dimensional vector embeddings for semantic similarity search.', style='List Bullet')
doc.add_paragraph('Goals: Convert text to numerical representations for similarity matching.', style='List Bullet')
doc.add_paragraph('Interface: Local Python library', style='List Bullet')

# 3.2 Use Cases by Package
add_heading_with_bookmark(doc, '3.2 Use Cases by Package', level=2)

# 3.2.1 Authentication Use Cases
add_heading_with_bookmark(doc, '3.2.1 Authentication Package', level=3)

doc.add_paragraph('UC-AUTH-001: Register Account')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows a new user to create an account by providing a valid email address.', style='List Bullet')
doc.add_paragraph('Preconditions: User is not logged in, email is not already registered.', style='List Bullet')
doc.add_paragraph('Postconditions: New user account created with unique UUID, user is logged in.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User navigates to registration page', style='List Bullet')
doc.add_paragraph('2. User enters email address', style='List Bullet')
doc.add_paragraph('3. System validates email format', style='List Bullet')
doc.add_paragraph('4. System checks email availability', style='List Bullet')
doc.add_paragraph('5. System creates new user record', style='List Bullet')
doc.add_paragraph('6. System logs user in automatically', style='List Bullet')
doc.add_paragraph('7. System redirects to home page', style='List Bullet')
doc.add_paragraph('Includes: Validate Email Format, Check Email Availability', style='List Bullet')

doc.add_paragraph('UC-AUTH-002: Login')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows a registered user to authenticate and access the application.', style='List Bullet')
doc.add_paragraph('Preconditions: User has a registered account, user is not logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: User is authenticated, session is created.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User navigates to login page', style='List Bullet')
doc.add_paragraph('2. User enters registered email address', style='List Bullet')
doc.add_paragraph('3. System validates email exists in database', style='List Bullet')
doc.add_paragraph('4. System creates user session', style='List Bullet')
doc.add_paragraph('5. System stores user ID in local storage', style='List Bullet')
doc.add_paragraph('6. System redirects to home page', style='List Bullet')
doc.add_paragraph('Includes: Maintain Session', style='List Bullet')

doc.add_paragraph('UC-AUTH-003: Logout')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows an authenticated user to end their session and log out.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: User session is terminated, user is redirected to login page.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User clicks logout button', style='List Bullet')
doc.add_paragraph('2. System clears session from local storage', style='List Bullet')
doc.add_paragraph('3. System redirects to login page', style='List Bullet')

doc.add_paragraph('UC-AUTH-004: Maintain Session')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Maintains user authentication state across page refreshes and browser tabs.', style='List Bullet')
doc.add_paragraph('Preconditions: User has previously logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: User remains authenticated if session is valid.', style='List Bullet')

# 3.2.2 Question Management Use Cases
add_heading_with_bookmark(doc, '3.2.2 Question Management Package', level=3)

doc.add_paragraph('UC-Q-001: Create Question')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows user to create a new question with optional answer.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: New question is saved to database, question appears in list.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User clicks "Add Question" button', style='List Bullet')
doc.add_paragraph('2. System opens modal dialog', style='List Bullet')
doc.add_paragraph('3. User enters question text (required)', style='List Bullet')
doc.add_paragraph('4. User optionally enters answer text', style='List Bullet')
doc.add_paragraph('5. User clicks "Save" button', style='List Bullet')
doc.add_paragraph('6. System validates question is not empty', style='List Bullet')
doc.add_paragraph('7. System saves question with user ID and timestamp', style='List Bullet')
doc.add_paragraph('8. System closes modal and refreshes list', style='List Bullet')
doc.add_paragraph('9. System displays success toast', style='List Bullet')
doc.add_paragraph('Includes: Open Modal Dialog, Validate Question Input, Display Toast Notification', style='List Bullet')

doc.add_paragraph('UC-Q-002: View Questions List')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Displays all questions belonging to the authenticated user.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: User\'s questions are displayed in sorted order.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User navigates to home page', style='List Bullet')
doc.add_paragraph('2. System fetches questions for current user', style='List Bullet')
doc.add_paragraph('3. System sorts questions by creation date (newest first)', style='List Bullet')
doc.add_paragraph('4. System displays questions with answer previews', style='List Bullet')
doc.add_paragraph('Includes: Display Empty State (when no questions exist)', style='List Bullet')

doc.add_paragraph('UC-Q-003: Edit Question')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows user to modify an existing question and/or answer.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in, question exists, question belongs to user.', style='List Bullet')
doc.add_paragraph('Postconditions: Question is updated with new values and timestamp.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User clicks edit option from question menu', style='List Bullet')
doc.add_paragraph('2. System opens modal with pre-populated data', style='List Bullet')
doc.add_paragraph('3. User modifies question text and/or answer', style='List Bullet')
doc.add_paragraph('4. User clicks "Save" button', style='List Bullet')
doc.add_paragraph('5. System validates and updates question', style='List Bullet')
doc.add_paragraph('6. System closes modal and refreshes list', style='List Bullet')
doc.add_paragraph('Includes: Open Modal Dialog, Validate Question Input, Display Toast Notification', style='List Bullet')

doc.add_paragraph('UC-Q-004: Delete Question')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows user to permanently remove a question.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in, question exists, question belongs to user.', style='List Bullet')
doc.add_paragraph('Postconditions: Question and associated answer are permanently deleted.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User clicks delete option from question menu', style='List Bullet')
doc.add_paragraph('2. System displays confirmation dialog', style='List Bullet')
doc.add_paragraph('3. User confirms deletion', style='List Bullet')
doc.add_paragraph('4. System deletes question from database', style='List Bullet')
doc.add_paragraph('5. System removes question from list', style='List Bullet')
doc.add_paragraph('6. System displays success toast', style='List Bullet')
doc.add_paragraph('Includes: Confirm Deletion, Display Toast Notification', style='List Bullet')

# 3.2.3 Answer Management Use Cases
add_heading_with_bookmark(doc, '3.2.3 Answer Management Package', level=3)

doc.add_paragraph('UC-A-001: Enter Manual Answer')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows user to manually type an answer to a question.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in, question form is open.', style='List Bullet')
doc.add_paragraph('Postconditions: Answer is saved with the question.', style='List Bullet')

doc.add_paragraph('UC-A-002: Request AI Answer')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Requests an AI-generated answer using the RAG pipeline.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in, question text is entered, documents are indexed.', style='List Bullet')
doc.add_paragraph('Postconditions: AI-generated answer is displayed in answer field.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. User enters question text', style='List Bullet')
doc.add_paragraph('2. User clicks "Ask Documents" button', style='List Bullet')
doc.add_paragraph('3. System displays loading indicator', style='List Bullet')
doc.add_paragraph('4. System triggers RAG pipeline', style='List Bullet')
doc.add_paragraph('5. System receives generated answer', style='List Bullet')
doc.add_paragraph('6. System populates answer in textarea', style='List Bullet')
doc.add_paragraph('Includes: Generate Answer from Documents, Show Loading Indicator', style='List Bullet')
doc.add_paragraph('Extends: Display Error Message (on failure)', style='List Bullet')

doc.add_paragraph('UC-A-003: Edit Answer')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows user to modify an existing answer (manual or AI-generated).', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in, answer exists.', style='List Bullet')
doc.add_paragraph('Postconditions: Answer is updated.', style='List Bullet')

doc.add_paragraph('UC-A-004: Regenerate AI Answer')
doc.add_paragraph('Actor: User', style='List Bullet')
doc.add_paragraph('Description: Allows user to request a new AI-generated answer.', style='List Bullet')
doc.add_paragraph('Preconditions: User is logged in, question text is entered.', style='List Bullet')
doc.add_paragraph('Postconditions: New answer replaces previous content in textarea.', style='List Bullet')
doc.add_paragraph('Includes: Generate Answer from Documents', style='List Bullet')

# 3.2.4 RAG Pipeline Use Cases
add_heading_with_bookmark(doc, '3.2.4 RAG Pipeline Package', level=3)

doc.add_paragraph('UC-RAG-001: Generate Answer from Documents')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Core RAG process that generates answers from indexed documents.', style='List Bullet')
doc.add_paragraph('Preconditions: Documents are indexed, question is provided.', style='List Bullet')
doc.add_paragraph('Postconditions: Generated answer is returned.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. System creates embedding for user question', style='List Bullet')
doc.add_paragraph('2. System searches for similar document chunks', style='List Bullet')
doc.add_paragraph('3. System retrieves top 3 matching chunks', style='List Bullet')
doc.add_paragraph('4. System assembles context from chunks', style='List Bullet')
doc.add_paragraph('5. System sends question and context to Claude API', style='List Bullet')
doc.add_paragraph('6. Claude generates natural language answer', style='List Bullet')
doc.add_paragraph('7. System returns answer to caller', style='List Bullet')
doc.add_paragraph('Includes: Create Query Embedding, Search Similar Chunks, Assemble Context, Generate LLM Response', style='List Bullet')
doc.add_paragraph('Extends: Handle Generation Error (on failure)', style='List Bullet')

doc.add_paragraph('UC-RAG-002: Search Similar Chunks')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Performs vector similarity search on indexed document chunks.', style='List Bullet')
doc.add_paragraph('Preconditions: Query embedding is created, chunks are indexed.', style='List Bullet')
doc.add_paragraph('Postconditions: Top matching chunks are returned with similarity scores.', style='List Bullet')

doc.add_paragraph('UC-RAG-003: Create Query Embedding')
doc.add_paragraph('Actor: Embedding Model', style='List Bullet')
doc.add_paragraph('Description: Converts question text to 384-dimensional vector embedding.', style='List Bullet')
doc.add_paragraph('Preconditions: Question text is provided.', style='List Bullet')
doc.add_paragraph('Postconditions: Embedding vector is returned.', style='List Bullet')

doc.add_paragraph('UC-RAG-004: Generate LLM Response')
doc.add_paragraph('Actor: Claude API', style='List Bullet')
doc.add_paragraph('Description: Generates natural language answer using Claude Haiku.', style='List Bullet')
doc.add_paragraph('Preconditions: Question and context are provided.', style='List Bullet')
doc.add_paragraph('Postconditions: Generated answer is returned.', style='List Bullet')

# 3.2.5 Document Management Use Cases
add_heading_with_bookmark(doc, '3.2.5 Document Management Package', level=3)

doc.add_paragraph('UC-DOC-001: Upload PDF Document')
doc.add_paragraph('Actor: Administrator', style='List Bullet')
doc.add_paragraph('Description: Uploads a PDF file to the server for indexing.', style='List Bullet')
doc.add_paragraph('Preconditions: Admin is logged in, file is valid PDF.', style='List Bullet')
doc.add_paragraph('Postconditions: PDF is stored in books folder.', style='List Bullet')
doc.add_paragraph('Includes: Detect Duplicate Document', style='List Bullet')

doc.add_paragraph('UC-DOC-002: Index Document')
doc.add_paragraph('Actor: Administrator', style='List Bullet')
doc.add_paragraph('Description: Processes and indexes a PDF document for RAG search.', style='List Bullet')
doc.add_paragraph('Preconditions: PDF is uploaded, admin is logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: Document is indexed with embeddings stored.', style='List Bullet')
doc.add_paragraph('Main Flow:', style='List Bullet')
doc.add_paragraph('1. Admin triggers indexing', style='List Bullet')
doc.add_paragraph('2. System extracts text from PDF', style='List Bullet')
doc.add_paragraph('3. System splits text into chunks (500 chars, 50 overlap)', style='List Bullet')
doc.add_paragraph('4. System generates embedding for each chunk', style='List Bullet')
doc.add_paragraph('5. System stores chunks and embeddings in database', style='List Bullet')
doc.add_paragraph('6. System records document metadata', style='List Bullet')
doc.add_paragraph('Includes: Extract Text, Split Chunks, Generate Embeddings, Store Metadata', style='List Bullet')

doc.add_paragraph('UC-DOC-003: View Indexed Documents')
doc.add_paragraph('Actor: Administrator', style='List Bullet')
doc.add_paragraph('Description: Displays list of all indexed documents with metadata.', style='List Bullet')
doc.add_paragraph('Preconditions: Admin is logged in.', style='List Bullet')
doc.add_paragraph('Postconditions: Document list is displayed.', style='List Bullet')

doc.add_paragraph('UC-DOC-004: Delete Indexed Document')
doc.add_paragraph('Actor: Administrator', style='List Bullet')
doc.add_paragraph('Description: Removes a document and its chunks from the index.', style='List Bullet')
doc.add_paragraph('Preconditions: Admin is logged in, document exists.', style='List Bullet')
doc.add_paragraph('Postconditions: Document and all associated chunks are deleted.', style='List Bullet')

# 3.2.6 User Interface Use Cases
add_heading_with_bookmark(doc, '3.2.6 User Interface Package', level=3)

doc.add_paragraph('UC-UI-001: Open Modal Dialog')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Opens a modal dialog for user interaction with focus trapping.', style='List Bullet')

doc.add_paragraph('UC-UI-002: Close Modal Dialog')
doc.add_paragraph('Actor: User/System', style='List Bullet')
doc.add_paragraph('Description: Closes modal via X button, backdrop click, or Escape key.', style='List Bullet')

doc.add_paragraph('UC-UI-003: Display Toast Notification')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Shows temporary success or error notification that auto-dismisses.', style='List Bullet')

doc.add_paragraph('UC-UI-004: Show Loading Indicator')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Displays loading state and disables interactive elements during async operations.', style='List Bullet')

doc.add_paragraph('UC-UI-005: Display Error Message')
doc.add_paragraph('Actor: System', style='List Bullet')
doc.add_paragraph('Description: Shows error feedback via toast or inline message.', style='List Bullet')

# 3.3 Relationships Summary
add_heading_with_bookmark(doc, '3.3 Relationships Summary', level=2)

doc.add_paragraph('Include Relationships (required behavior):')
doc.add_paragraph('Register Account includes Validate Email Format, Check Email Availability', style='List Bullet')
doc.add_paragraph('Login includes Maintain Session', style='List Bullet')
doc.add_paragraph('Create Question includes Open Modal, Validate Question, Display Toast', style='List Bullet')
doc.add_paragraph('Edit Question includes Open Modal, Validate Question, Display Toast', style='List Bullet')
doc.add_paragraph('Delete Question includes Confirm Deletion, Display Toast', style='List Bullet')
doc.add_paragraph('Request AI Answer includes Generate Answer from Documents, Show Loading', style='List Bullet')
doc.add_paragraph('Generate Answer includes Query Embedding, Search Chunks, Assemble Context, LLM Response', style='List Bullet')
doc.add_paragraph('Index Document includes Extract Text, Split Chunks, Generate Embeddings, Store Metadata', style='List Bullet')
doc.add_paragraph('Upload PDF includes Detect Duplicate', style='List Bullet')

doc.add_paragraph('Extend Relationships (optional behavior):')
doc.add_paragraph('Handle Generation Error extends Generate Answer from Documents', style='List Bullet')
doc.add_paragraph('Display Error Message extends Request AI Answer', style='List Bullet')

doc.add_paragraph('Generalization Relationships (inheritance):')
doc.add_paragraph('Administrator inherits from User (has all user capabilities plus admin functions)', style='List Bullet')

# 3.4 PlantUML Diagram Code
add_heading_with_bookmark(doc, '3.4 PlantUML Diagram Code', level=2)
doc.add_paragraph(
    'The following PlantUML code can be used to generate a visual representation of the '
    'Use Case Diagram. This code can be rendered using PlantUML online editor, VS Code '
    'PlantUML extension, or any compatible tool.'
)

doc.add_paragraph('File location: diagrams/use_case_diagram.puml')

# Add Use Case Diagram Image
doc.add_picture('diagrams/use_case_diagram.png', width=Inches(6))

# Figure Description
fig_para = doc.add_paragraph()
fig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run = fig_para.add_run('Figure 1: Use Case Diagram for Ask Holmes Application')
fig_run.italic = True

# ============================================
# 4. ACTIVITY DIAGRAM - QUESTIONS CRUD
# ============================================
add_heading_with_bookmark(doc, '4. Activity Diagram - Questions CRUD Operations', level=1)

doc.add_paragraph(
    'This chapter presents the activity diagrams for the basic CRUD (Create, Read, Update, Delete) '
    'operations on questions within the Ask Holmes application. These diagrams illustrate the flow '
    'of activities between the User and the System for each operation.'
)

# 4.1 Create Question
add_heading_with_bookmark(doc, '4.1 Create Question', level=2)

doc.add_paragraph('Actors: User, System')
doc.add_paragraph('Precondition: User is authenticated and on the Questions page')
doc.add_paragraph('Postcondition: New question is created and displayed in the list')

doc.add_paragraph('Flow of Activities:')
doc.add_paragraph('1. User clicks "New Question" button', style='List Number')
doc.add_paragraph('2. System opens Create Question Modal with empty form', style='List Number')
doc.add_paragraph('3. User enters question title', style='List Number')
doc.add_paragraph('4. User enters question content', style='List Number')
doc.add_paragraph('5. User clicks "Save" button', style='List Number')
doc.add_paragraph('6. System validates input', style='List Number')
doc.add_paragraph('7. If valid: System sends POST request to /api/questions', style='List Number')
doc.add_paragraph('8. System creates question in database with timestamp', style='List Number')
doc.add_paragraph('9. System closes modal and displays success toast', style='List Number')
doc.add_paragraph('10. System refreshes questions list', style='List Number')

doc.add_paragraph('Alternative Flow (Invalid Input):')
doc.add_paragraph('6a. System displays validation error message', style='List Bullet')
doc.add_paragraph('6b. System highlights invalid fields', style='List Bullet')
doc.add_paragraph('6c. User corrects input errors and returns to step 5', style='List Bullet')

# 4.2 Read Questions (View List)
add_heading_with_bookmark(doc, '4.2 Read Questions (View List)', level=2)

doc.add_paragraph('Actors: User, System')
doc.add_paragraph('Precondition: User is authenticated')
doc.add_paragraph('Postcondition: Questions list is displayed')

doc.add_paragraph('Flow of Activities:')
doc.add_paragraph('1. User navigates to Questions page', style='List Number')
doc.add_paragraph('2. System sends GET request to /api/questions', style='List Number')
doc.add_paragraph('3. System fetches user questions from database', style='List Number')
doc.add_paragraph('4. If questions exist: System renders questions list with cards', style='List Number')
doc.add_paragraph('5. Each card displays title, preview, and date', style='List Number')

doc.add_paragraph('Alternative Flow (No Questions):')
doc.add_paragraph('4a. System displays empty state message', style='List Bullet')
doc.add_paragraph('4b. System shows "Create your first question" prompt', style='List Bullet')

# 4.3 Read Question (View Details)
add_heading_with_bookmark(doc, '4.3 Read Question (View Details)', level=2)

doc.add_paragraph('Actors: User, System')
doc.add_paragraph('Precondition: User is on Questions list page')
doc.add_paragraph('Postcondition: Question details are displayed with answer section')

doc.add_paragraph('Flow of Activities:')
doc.add_paragraph('1. User clicks on question card', style='List Number')
doc.add_paragraph('2. System sends GET request to /api/questions/{id}', style='List Number')
doc.add_paragraph('3. System fetches question details and associated answer', style='List Number')
doc.add_paragraph('4. System navigates to Question Details page', style='List Number')
doc.add_paragraph('5. System displays question title and content', style='List Number')
doc.add_paragraph('6. If answer exists: System displays answer with Edit/Clear buttons', style='List Number')

doc.add_paragraph('Alternative Flow (No Answer):')
doc.add_paragraph('6a. System displays empty answer section', style='List Bullet')
doc.add_paragraph('6b. System shows "Enter Answer" and "Ask Holmes" buttons', style='List Bullet')

# 4.4 Update Question
add_heading_with_bookmark(doc, '4.4 Update Question', level=2)

doc.add_paragraph('Actors: User, System')
doc.add_paragraph('Precondition: User is viewing a question they own')
doc.add_paragraph('Postcondition: Question is updated with new content')

doc.add_paragraph('Flow of Activities:')
doc.add_paragraph('1. User clicks "Edit" button on question', style='List Number')
doc.add_paragraph('2. System opens Edit Question Modal', style='List Number')
doc.add_paragraph('3. System populates form with existing question data', style='List Number')
doc.add_paragraph('4. User modifies question title and/or content', style='List Number')
doc.add_paragraph('5. User clicks "Save Changes" button', style='List Number')
doc.add_paragraph('6. System validates input', style='List Number')
doc.add_paragraph('7. If valid: System sends PUT request to /api/questions/{id}', style='List Number')
doc.add_paragraph('8. System updates question and modified timestamp', style='List Number')
doc.add_paragraph('9. System closes modal and displays success toast', style='List Number')
doc.add_paragraph('10. System refreshes question display', style='List Number')

doc.add_paragraph('Alternative Flow (Invalid Input):')
doc.add_paragraph('6a. System displays validation error message', style='List Bullet')
doc.add_paragraph('6b. System highlights invalid fields', style='List Bullet')
doc.add_paragraph('6c. User corrects input errors and returns to step 5', style='List Bullet')

# 4.5 Delete Question
add_heading_with_bookmark(doc, '4.5 Delete Question', level=2)

doc.add_paragraph('Actors: User, System')
doc.add_paragraph('Precondition: User is viewing a question they own')
doc.add_paragraph('Postcondition: Question and associated answer are deleted')

doc.add_paragraph('Flow of Activities:')
doc.add_paragraph('1. User clicks "Delete" button on question', style='List Number')
doc.add_paragraph('2. System opens confirmation dialog', style='List Number')
doc.add_paragraph('3. System displays "Are you sure?" message', style='List Number')
doc.add_paragraph('4. User clicks "Confirm Delete" button', style='List Number')
doc.add_paragraph('5. System sends DELETE request to /api/questions/{id}', style='List Number')
doc.add_paragraph('6. System deletes question and associated answer from database', style='List Number')
doc.add_paragraph('7. System closes dialog and displays success toast', style='List Number')
doc.add_paragraph('8. System removes question from list', style='List Number')
doc.add_paragraph('9. System navigates to Questions list page', style='List Number')

doc.add_paragraph('Alternative Flow (Cancel Deletion):')
doc.add_paragraph('4a. User clicks "Cancel" button', style='List Bullet')
doc.add_paragraph('4b. System closes confirmation dialog', style='List Bullet')
doc.add_paragraph('4c. System returns to previous state', style='List Bullet')

# 4.6 PlantUML Diagram Code
add_heading_with_bookmark(doc, '4.6 PlantUML Diagram Code', level=2)
doc.add_paragraph(
    'The following PlantUML code can be used to generate a visual representation of the '
    'Activity Diagram. This code can be rendered using PlantUML online editor, VS Code '
    'PlantUML extension, or any compatible tool.'
)

doc.add_paragraph('File location: diagrams/activity_diagram_questions_crud.puml')

# Add Activity Diagram Image
doc.add_picture('diagrams/activity_diagram.png', width=Inches(6))

# Figure Description
fig_para2 = doc.add_paragraph()
fig_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run2 = fig_para2.add_run('Figure 2: Activity Diagram for Questions CRUD Operations')
fig_run2.italic = True

# ============================================
# 5. SEQUENCE DIAGRAM - QUESTIONS CRUD
# ============================================
add_heading_with_bookmark(doc, '5. Sequence Diagram - Questions CRUD Operations', level=1)

doc.add_paragraph(
    'This chapter presents the sequence diagrams for the basic CRUD (Create, Read, Update, Delete) '
    'operations on questions within the Ask Holmes application. These diagrams illustrate the '
    'interactions between the User, React Frontend, Flask API Backend, and PostgreSQL Database '
    'over time for each operation.'
)

# 5.1 Participants
add_heading_with_bookmark(doc, '5.1 Participants', level=2)

doc.add_paragraph('User: The actor initiating all operations', style='List Bullet')
doc.add_paragraph('React Frontend: The client-side user interface handling user interactions and state management', style='List Bullet')
doc.add_paragraph('Flask API: The backend server processing requests and business logic', style='List Bullet')
doc.add_paragraph('PostgreSQL: The database storing questions and answers data', style='List Bullet')

# 5.2 Create Question Sequence
add_heading_with_bookmark(doc, '5.2 Create Question Sequence', level=2)

doc.add_paragraph('1. User clicks New Question button', style='List Number')
doc.add_paragraph('2. Frontend opens CreateQuestionModal and displays empty form', style='List Number')
doc.add_paragraph('3. User enters title and content, clicks Save', style='List Number')
doc.add_paragraph('4. Frontend validates input', style='List Number')
doc.add_paragraph('5. Frontend sends POST /api/questions to Backend', style='List Number')
doc.add_paragraph('6. Backend validates request and gets current user from session', style='List Number')
doc.add_paragraph('7. Backend inserts question into Database', style='List Number')
doc.add_paragraph('8. Database returns new question ID', style='List Number')
doc.add_paragraph('9. Backend returns 201 Created to Frontend', style='List Number')
doc.add_paragraph('10. Frontend closes modal, updates state, displays success toast', style='List Number')

# 5.3 Read Questions Sequence
add_heading_with_bookmark(doc, '5.3 Read Questions (List) Sequence', level=2)

doc.add_paragraph('1. User navigates to Questions page', style='List Number')
doc.add_paragraph('2. Frontend sends GET /api/questions to Backend', style='List Number')
doc.add_paragraph('3. Backend gets current user from session', style='List Number')
doc.add_paragraph('4. Backend queries Database for user questions', style='List Number')
doc.add_paragraph('5. Database returns questions array', style='List Number')
doc.add_paragraph('6. Backend returns 200 OK with questions list', style='List Number')
doc.add_paragraph('7. Frontend renders questions list or empty state', style='List Number')

# 5.4 Read Question Details Sequence
add_heading_with_bookmark(doc, '5.4 Read Question (Details) Sequence', level=2)

doc.add_paragraph('1. User clicks on question card', style='List Number')
doc.add_paragraph('2. Frontend sends GET /api/questions/id to Backend', style='List Number')
doc.add_paragraph('3. Backend gets current user and queries question from Database', style='List Number')
doc.add_paragraph('4. Backend queries associated answer from Database', style='List Number')
doc.add_paragraph('5. Backend returns 200 OK with question and answer data', style='List Number')
doc.add_paragraph('6. Frontend navigates to details page and displays content', style='List Number')

# 5.5 Update Question Sequence
add_heading_with_bookmark(doc, '5.5 Update Question Sequence', level=2)

doc.add_paragraph('1. User clicks Edit button', style='List Number')
doc.add_paragraph('2. Frontend opens EditQuestionModal with existing data', style='List Number')
doc.add_paragraph('3. User modifies title/content and clicks Save Changes', style='List Number')
doc.add_paragraph('4. Frontend validates input', style='List Number')
doc.add_paragraph('5. Frontend sends PUT /api/questions/id to Backend', style='List Number')
doc.add_paragraph('6. Backend validates and verifies ownership in Database', style='List Number')
doc.add_paragraph('7. Backend updates question in Database', style='List Number')
doc.add_paragraph('8. Backend returns 200 OK with updated question', style='List Number')
doc.add_paragraph('9. Frontend closes modal, updates state, displays success toast', style='List Number')

# 5.6 Delete Question Sequence
add_heading_with_bookmark(doc, '5.6 Delete Question Sequence', level=2)

doc.add_paragraph('1. User clicks Delete button', style='List Number')
doc.add_paragraph('2. Frontend opens confirmation dialog', style='List Number')
doc.add_paragraph('3. User clicks Confirm Delete', style='List Number')
doc.add_paragraph('4. Frontend sends DELETE /api/questions/id to Backend', style='List Number')
doc.add_paragraph('5. Backend verifies ownership in Database', style='List Number')
doc.add_paragraph('6. Backend deletes associated answer from Database', style='List Number')
doc.add_paragraph('7. Backend deletes question from Database', style='List Number')
doc.add_paragraph('8. Backend returns 200 OK', style='List Number')
doc.add_paragraph('9. Frontend closes dialog, removes from state, displays success toast', style='List Number')
doc.add_paragraph('10. Frontend navigates to questions list', style='List Number')

# 5.7 Diagram
add_heading_with_bookmark(doc, '5.7 Sequence Diagram', level=2)

doc.add_paragraph('File location: diagrams/sequence_diagram_questions_crud.puml')

# Add Sequence Diagram Image
doc.add_picture('diagrams/sequence_diagram.png', width=Inches(4.2))

# Figure Description
fig_para3 = doc.add_paragraph()
fig_para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run3 = fig_para3.add_run('Figure 3: Sequence Diagram for Questions CRUD Operations')
fig_run3.italic = True

# ============================================
# 6. CLASS DIAGRAM
# ============================================
add_heading_with_bookmark(doc, '6. Class Diagram', level=1)

doc.add_paragraph(
    'This chapter presents a comprehensive class diagram for the Ask Holmes application, '
    'illustrating the object-oriented structure of both the backend (Python/Flask) and '
    'frontend (React) components. The diagram shows classes, their attributes, methods, '
    'and relationships across all layers of the application.'
)

# 6.1 Overview
add_heading_with_bookmark(doc, '6.1 Overview', level=2)

doc.add_paragraph(
    'The Ask Holmes application follows a three-tier architecture with clear separation '
    'between the presentation layer (React frontend), business logic layer (Flask services), '
    'and data access layer (SQLAlchemy models with PostgreSQL). The class diagram is organized '
    'into the following packages:'
)

doc.add_paragraph('Backend Models: SQLAlchemy ORM entities representing database tables (User, Question, Document, DocumentChunk)', style='List Bullet')
doc.add_paragraph('Backend Services: Business logic classes handling operations (AuthService, QuestionService, RAGService)', style='List Bullet')
doc.add_paragraph('Backend Utilities: Helper classes for embeddings, PDF processing, and LLM interaction (Config, EmbeddingModel, PDFProcessor, LLMClient)', style='List Bullet')
doc.add_paragraph('Backend API Routes: Flask Blueprint route handlers (questions_bp, auth_bp, rag_bp)', style='List Bullet')
doc.add_paragraph('Frontend Pages: React page components (LoginPage, HomePage)', style='List Bullet')
doc.add_paragraph('Frontend Components: Reusable React UI components (Header, Modal, Toast, QuestionList, QuestionItem, QuestionForm)', style='List Bullet')
doc.add_paragraph('Frontend Services: API client modules for backend communication (authApi, questionsApi, ragApi)', style='List Bullet')
doc.add_paragraph('Frontend Context: React Context for state management (AuthContext, App)', style='List Bullet')
doc.add_paragraph('External Systems: PostgreSQL database, Claude API, and Sentence Transformers', style='List Bullet')

# 6.2 Backend Models
add_heading_with_bookmark(doc, '6.2 Backend Models', level=2)

doc.add_paragraph(
    'The backend model layer consists of four SQLAlchemy ORM classes that map to PostgreSQL '
    'database tables. All models implement a to_dict() method for JSON serialization.'
)

doc.add_paragraph('User')
doc.add_paragraph('Purpose: Represents registered users of the application', style='List Bullet')
doc.add_paragraph('Attributes: id (UUID, primary key), email (unique, indexed), created_at (timestamp)', style='List Bullet')
doc.add_paragraph('Relationships: One-to-many with Question (cascade delete)', style='List Bullet')

doc.add_paragraph('Question')
doc.add_paragraph('Purpose: Stores user questions and their answers', style='List Bullet')
doc.add_paragraph('Attributes: id (UUID), user_id (foreign key), question (text), answer (text), created_at, updated_at', style='List Bullet')
doc.add_paragraph('Relationships: Many-to-one with User', style='List Bullet')

doc.add_paragraph('Document')
doc.add_paragraph('Purpose: Tracks indexed PDF documents in the RAG system', style='List Bullet')
doc.add_paragraph('Attributes: id (UUID), filename, title, file_hash (unique), indexed_at, chunk_count', style='List Bullet')
doc.add_paragraph('Relationships: One-to-many with DocumentChunk (cascade delete)', style='List Bullet')

doc.add_paragraph('DocumentChunk')
doc.add_paragraph('Purpose: Stores vectorized text segments for semantic search', style='List Bullet')
doc.add_paragraph('Attributes: id (UUID), document_id (foreign key), chunk_text, chunk_index, embedding (384-dim vector), created_at', style='List Bullet')
doc.add_paragraph('Relationships: Many-to-one with Document', style='List Bullet')

# 6.3 Backend Services
add_heading_with_bookmark(doc, '6.3 Backend Services', level=2)

doc.add_paragraph(
    'The service layer implements the business logic and acts as an intermediary between '
    'the API routes and the data models. Services use static methods to maintain statelessness.'
)

doc.add_paragraph('AuthService')
doc.add_paragraph('Purpose: Handles user authentication and registration', style='List Bullet')
doc.add_paragraph('Key Methods: authenticate_by_email(), get_user_by_id(), create_user()', style='List Bullet')

doc.add_paragraph('QuestionService')
doc.add_paragraph('Purpose: Manages CRUD operations for user questions', style='List Bullet')
doc.add_paragraph('Key Methods: get_user_questions(), get_question_by_id(), create_question(), update_question(), delete_question()', style='List Bullet')

doc.add_paragraph('RAGService')
doc.add_paragraph('Purpose: Orchestrates the Retrieval-Augmented Generation pipeline', style='List Bullet')
doc.add_paragraph('Dependencies: PDFProcessor, EmbeddingModel, LLMClient', style='List Bullet')
doc.add_paragraph('Key Methods: index_document(), search_similar_chunks(), generate_answer(), get_indexed_documents(), delete_document()', style='List Bullet')

# 6.4 Backend Utilities
add_heading_with_bookmark(doc, '6.4 Backend Utilities', level=2)

doc.add_paragraph('Config: Centralized application configuration using environment variables (DATABASE_URL, ANTHROPIC_API_KEY, CHUNK_SIZE=500, CHUNK_OVERLAP=50, TOP_K_RESULTS=3, EMBEDDING_DIMENSION=384)', style='List Bullet')
doc.add_paragraph('EmbeddingModel (Singleton): Generates 384-dimensional vector embeddings using Sentence Transformers (all-MiniLM-L6-v2 model)', style='List Bullet')
doc.add_paragraph('PDFProcessor: Extracts text from PDFs, splits into overlapping chunks, calculates file hashes for duplicate detection', style='List Bullet')
doc.add_paragraph('LLMClient: Interfaces with Anthropic Claude API (claude-3-haiku) for answer generation', style='List Bullet')

# 6.5 Backend API Routes
add_heading_with_bookmark(doc, '6.5 Backend API Routes', level=2)

doc.add_paragraph('questions_bp: /api/questions - GET (list), GET/<id> (detail), POST (create), PUT/<id> (update), DELETE/<id> (delete)', style='List Bullet')
doc.add_paragraph('auth_bp: /api/auth - POST /login, POST /register, GET /user/<id>', style='List Bullet')
doc.add_paragraph('rag_bp: /api/rag - POST /index, POST /query, POST /search, GET /documents, DELETE /documents/<id>', style='List Bullet')

# 6.6 Frontend Components
add_heading_with_bookmark(doc, '6.6 Frontend Components', level=2)

doc.add_paragraph('Pages:')
doc.add_paragraph('LoginPage: Authentication entry point with email input, loading states, and error handling', style='List Bullet')
doc.add_paragraph('HomePage: Main view for managing questions with CRUD operations and RAG integration', style='List Bullet')

doc.add_paragraph('UI Components:')
doc.add_paragraph('Header: Application header with branding and logout functionality', style='List Bullet')
doc.add_paragraph('Modal: Reusable dialog wrapper with escape key and backdrop click handling', style='List Bullet')
doc.add_paragraph('Toast: Temporary notification messages (success/error)', style='List Bullet')
doc.add_paragraph('QuestionList: Renders list of questions or empty state', style='List Bullet')
doc.add_paragraph('QuestionItem: Individual question card with actions menu and two-step delete confirmation', style='List Bullet')
doc.add_paragraph('QuestionForm: Form for creating/editing questions with manual entry and AI generation', style='List Bullet')

# 6.7 Frontend Services and Context
add_heading_with_bookmark(doc, '6.7 Frontend Services and Context', level=2)

doc.add_paragraph('API Services:')
doc.add_paragraph('authApi: login(), register(), getUser() - handles authentication requests', style='List Bullet')
doc.add_paragraph('questionsApi: getAll(), getOne(), create(), update(), delete() - manages question CRUD via HTTP', style='List Bullet')
doc.add_paragraph('ragApi: indexDocuments(), query(), getDocuments(), deleteDocument(), searchChunks() - RAG operations', style='List Bullet')

doc.add_paragraph('Context:')
doc.add_paragraph('AuthContext: Global authentication state (user, login, logout) with localStorage persistence', style='List Bullet')
doc.add_paragraph('App: Root component providing AuthContext, renders LoginPage or HomePage based on auth state', style='List Bullet')

# 6.8 External Systems
add_heading_with_bookmark(doc, '6.8 External Systems', level=2)

doc.add_paragraph('PostgreSQL Database: Persistent storage with pgvector extension for vector similarity search using IVFFlat indexing', style='List Bullet')
doc.add_paragraph('Anthropic Claude API: claude-3-haiku model for natural language answer generation', style='List Bullet')
doc.add_paragraph('Sentence Transformers: all-MiniLM-L6-v2 model for generating 384-dimensional text embeddings', style='List Bullet')

# 6.9 Key Relationships
add_heading_with_bookmark(doc, '6.9 Key Relationships', level=2)

doc.add_paragraph('Composition (Strong Ownership):')
doc.add_paragraph('User "1" *-- "0..*" Question (cascade delete)', style='List Bullet')
doc.add_paragraph('Document "1" *-- "0..*" DocumentChunk (cascade delete)', style='List Bullet')

doc.add_paragraph('Dependencies:')
doc.add_paragraph('RAGService uses PDFProcessor, EmbeddingModel, and LLMClient', style='List Bullet')
doc.add_paragraph('Services manage their respective Models', style='List Bullet')
doc.add_paragraph('Routes delegate to Services', style='List Bullet')
doc.add_paragraph('Frontend API modules communicate with Backend Routes via HTTP', style='List Bullet')
doc.add_paragraph('React components use AuthContext for authentication state', style='List Bullet')

# 6.10 Class Diagram
add_heading_with_bookmark(doc, '6.10 Class Diagram', level=2)

doc.add_paragraph('File location: diagrams/class_diagram.puml')

# Add Class Diagram Image
doc.add_picture('diagrams/class_diagram.png', width=Inches(6.5))

# Figure Description
fig_para4 = doc.add_paragraph()
fig_para4.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run4 = fig_para4.add_run('Figure 4: Class Diagram for Ask Holmes Application')
fig_run4.italic = True

# ============================================
# 7. ENTITY RELATIONSHIP DIAGRAM
# ============================================
add_heading_with_bookmark(doc, '7. Entity Relationship Diagram', level=1)

doc.add_paragraph(
    'This chapter presents the Entity Relationship Diagram (ERD) for the Ask Holmes application, '
    'illustrating the database schema and relationships between entities. The application uses '
    'PostgreSQL as its relational database management system with the pgvector extension for '
    'vector similarity search capabilities.'
)

# 7.1 Overview
add_heading_with_bookmark(doc, '7.1 Overview', level=2)

doc.add_paragraph(
    'The database schema consists of four main entities that support the core functionality '
    'of the Ask Holmes application: user management, question-answer storage, and document '
    'indexing for the RAG (Retrieval-Augmented Generation) pipeline.'
)

doc.add_paragraph('The schema follows these design principles:')
doc.add_paragraph('UUID Primary Keys: All entities use UUID as primary keys for global uniqueness and security', style='List Bullet')
doc.add_paragraph('Referential Integrity: Foreign key constraints ensure data consistency', style='List Bullet')
doc.add_paragraph('Cascade Deletion: Child records are automatically deleted when parent records are removed', style='List Bullet')
doc.add_paragraph('Indexing Strategy: Frequently queried columns are indexed for optimal performance', style='List Bullet')
doc.add_paragraph('Timestamp Tracking: All entities track creation time; mutable entities track update time', style='List Bullet')

# 7.2 Entities
add_heading_with_bookmark(doc, '7.2 Entities', level=2)

doc.add_paragraph('USERS')
doc.add_paragraph('Purpose: Stores registered user accounts for authentication', style='List Bullet')
doc.add_paragraph('Primary Key: id (UUID, auto-generated)', style='List Bullet')
doc.add_paragraph('Attributes:', style='List Bullet')
doc.add_paragraph('  - email: VARCHAR(255), unique, indexed, not null - User email address for login', style='List Bullet')
doc.add_paragraph('  - created_at: TIMESTAMP, default current timestamp - Account creation time', style='List Bullet')
doc.add_paragraph('Constraints: Unique constraint on email ensures no duplicate accounts', style='List Bullet')

doc.add_paragraph('QUESTIONS')
doc.add_paragraph('Purpose: Stores user questions and their associated answers', style='List Bullet')
doc.add_paragraph('Primary Key: id (UUID, auto-generated)', style='List Bullet')
doc.add_paragraph('Foreign Key: user_id references USERS(id) with CASCADE delete', style='List Bullet')
doc.add_paragraph('Attributes:', style='List Bullet')
doc.add_paragraph('  - user_id: UUID, indexed, not null - Owner of the question', style='List Bullet')
doc.add_paragraph('  - question: TEXT, not null - The question content', style='List Bullet')
doc.add_paragraph('  - answer: TEXT, nullable - The answer (manual or AI-generated)', style='List Bullet')
doc.add_paragraph('  - created_at: TIMESTAMP, indexed, default current timestamp - Question creation time', style='List Bullet')
doc.add_paragraph('  - updated_at: TIMESTAMP, auto-updated - Last modification time', style='List Bullet')
doc.add_paragraph('Indexes: user_id and created_at are indexed for efficient filtering and sorting', style='List Bullet')

doc.add_paragraph('DOCUMENTS')
doc.add_paragraph('Purpose: Tracks PDF documents that have been indexed for RAG search', style='List Bullet')
doc.add_paragraph('Primary Key: id (UUID, auto-generated)', style='List Bullet')
doc.add_paragraph('Attributes:', style='List Bullet')
doc.add_paragraph('  - filename: VARCHAR(255), indexed, not null - Original PDF filename', style='List Bullet')
doc.add_paragraph('  - title: VARCHAR(500), nullable - Document title (extracted from filename)', style='List Bullet')
doc.add_paragraph('  - file_hash: VARCHAR(64), unique, not null - SHA-256 hash for duplicate detection', style='List Bullet')
doc.add_paragraph('  - indexed_at: TIMESTAMP, default current timestamp - When document was indexed', style='List Bullet')
doc.add_paragraph('  - chunk_count: INTEGER, default 0 - Number of text chunks extracted', style='List Bullet')
doc.add_paragraph('Constraints: Unique constraint on file_hash prevents duplicate document indexing', style='List Bullet')

doc.add_paragraph('DOCUMENT_CHUNKS')
doc.add_paragraph('Purpose: Stores vectorized text segments for semantic similarity search', style='List Bullet')
doc.add_paragraph('Primary Key: id (UUID, auto-generated)', style='List Bullet')
doc.add_paragraph('Foreign Key: document_id references DOCUMENTS(id) with CASCADE delete', style='List Bullet')
doc.add_paragraph('Attributes:', style='List Bullet')
doc.add_paragraph('  - document_id: UUID, indexed, not null - Parent document reference', style='List Bullet')
doc.add_paragraph('  - chunk_text: TEXT, not null - The text content of the chunk', style='List Bullet')
doc.add_paragraph('  - chunk_index: INTEGER, not null - Position of chunk within document', style='List Bullet')
doc.add_paragraph('  - embedding: VECTOR(384), pgvector type - 384-dimensional embedding vector', style='List Bullet')
doc.add_paragraph('  - created_at: TIMESTAMP, default current timestamp - Chunk creation time', style='List Bullet')
doc.add_paragraph('Special: Uses pgvector extension with IVFFlat index for efficient similarity search', style='List Bullet')

# 7.3 Relationships
add_heading_with_bookmark(doc, '7.3 Relationships', level=2)

doc.add_paragraph('USERS to QUESTIONS (One-to-Many)')
doc.add_paragraph('Cardinality: One user can have zero or more questions', style='List Bullet')
doc.add_paragraph('Relationship: USERS ||--o{ QUESTIONS', style='List Bullet')
doc.add_paragraph('Foreign Key: QUESTIONS.user_id references USERS.id', style='List Bullet')
doc.add_paragraph('On Delete: CASCADE - Deleting a user removes all their questions', style='List Bullet')
doc.add_paragraph('Business Rule: Questions are isolated per user; users can only access their own questions', style='List Bullet')

doc.add_paragraph('DOCUMENTS to DOCUMENT_CHUNKS (One-to-Many)')
doc.add_paragraph('Cardinality: One document can have zero or more chunks', style='List Bullet')
doc.add_paragraph('Relationship: DOCUMENTS ||--o{ DOCUMENT_CHUNKS', style='List Bullet')
doc.add_paragraph('Foreign Key: DOCUMENT_CHUNKS.document_id references DOCUMENTS.id', style='List Bullet')
doc.add_paragraph('On Delete: CASCADE - Deleting a document removes all its chunks', style='List Bullet')
doc.add_paragraph('Business Rule: Chunks are created during indexing; typical document has 50-500 chunks', style='List Bullet')

# 7.4 Indexes and Performance
add_heading_with_bookmark(doc, '7.4 Indexes and Performance', level=2)

doc.add_paragraph('Standard B-tree Indexes:')
doc.add_paragraph('USERS.email - For fast login lookups', style='List Bullet')
doc.add_paragraph('QUESTIONS.user_id - For filtering questions by user', style='List Bullet')
doc.add_paragraph('QUESTIONS.created_at - For sorting questions by date', style='List Bullet')
doc.add_paragraph('DOCUMENTS.filename - For document name searches', style='List Bullet')
doc.add_paragraph('DOCUMENT_CHUNKS.document_id - For retrieving chunks by document', style='List Bullet')

doc.add_paragraph('Vector Index (pgvector):')
doc.add_paragraph('DOCUMENT_CHUNKS.embedding - IVFFlat index for approximate nearest neighbor search', style='List Bullet')
doc.add_paragraph('Configuration: Lists parameter tuned for dataset size (typically 100 lists for up to 100K chunks)', style='List Bullet')
doc.add_paragraph('Search Method: Cosine distance for semantic similarity matching', style='List Bullet')

# 7.5 Entity Relationship Diagram
add_heading_with_bookmark(doc, '7.5 Entity Relationship Diagram', level=2)

doc.add_paragraph('File location: diagrams/er_diagram.puml')

# Add ER Diagram Image
doc.add_picture('diagrams/er_diagram.png', width=Inches(6))

# Figure Description
fig_para5 = doc.add_paragraph()
fig_para5.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run5 = fig_para5.add_run('Figure 5: Entity Relationship Diagram for Ask Holmes Database')
fig_run5.italic = True

# ============================================
# 8. DESIGN PATTERNS
# ============================================
add_heading_with_bookmark(doc, '8. Design Patterns', level=1)

doc.add_paragraph(
    'This chapter documents the software design patterns employed in the Ask Holmes application. '
    'Design patterns provide proven solutions to common software design problems, improving code '
    'maintainability, reusability, and scalability. The application utilizes patterns across both '
    'the backend (Python/Flask) and frontend (React) layers.'
)

# 8.1 Backend Design Patterns
add_heading_with_bookmark(doc, '8.1 Backend Design Patterns', level=2)

doc.add_paragraph('8.1.1 Application Factory Pattern')
doc.add_paragraph('Location: backend/app/__init__.py', style='List Bullet')
doc.add_paragraph('Implementation: The create_app() function creates and configures the Flask application instance.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Enables multiple app instances with different configurations (testing, development, production)', style='List Bullet')
doc.add_paragraph('  - Delays extension initialization until app creation time', style='List Bullet')
doc.add_paragraph('  - Supports blueprint registration for modular architecture', style='List Bullet')

doc.add_paragraph('8.1.2 Singleton Pattern')
doc.add_paragraph('Location: backend/app/utils/embeddings.py (EmbeddingModel class)', style='List Bullet')
doc.add_paragraph('Implementation: Uses __new__ method and class-level variables to ensure only one instance of the embedding model is created.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Prevents multiple loading of the ML model into memory', style='List Bullet')
doc.add_paragraph('  - Reduces memory footprint significantly (model is ~90MB)', style='List Bullet')
doc.add_paragraph('  - Ensures consistent embedding generation across requests', style='List Bullet')

doc.add_paragraph('8.1.3 Service Layer Pattern')
doc.add_paragraph('Location: backend/app/services/ (AuthService, QuestionService, RAGService)', style='List Bullet')
doc.add_paragraph('Implementation: Service classes encapsulate business logic, separating it from route handlers and data models.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Clear separation of concerns between presentation and business logic', style='List Bullet')
doc.add_paragraph('  - Enables unit testing of business logic in isolation', style='List Bullet')
doc.add_paragraph('  - Promotes code reuse across different routes', style='List Bullet')

doc.add_paragraph('8.1.4 Repository Pattern')
doc.add_paragraph('Location: backend/app/services/ (implemented within service classes)', style='List Bullet')
doc.add_paragraph('Implementation: Services abstract database operations, providing a clean API for data access.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Abstracts data persistence mechanism from business logic', style='List Bullet')
doc.add_paragraph('  - Simplifies switching between different data sources', style='List Bullet')
doc.add_paragraph('  - Centralizes query logic for maintainability', style='List Bullet')

doc.add_paragraph('8.1.5 Blueprint Pattern')
doc.add_paragraph('Location: backend/app/routes/ (auth_bp, questions_bp, rag_bp)', style='List Bullet')
doc.add_paragraph('Implementation: Flask Blueprints organize routes into modular, reusable components.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Modular organization of related routes', style='List Bullet')
doc.add_paragraph('  - URL prefix management (/api/auth, /api/questions, /api/rag)', style='List Bullet')
doc.add_paragraph('  - Enables lazy loading and separate testing of route modules', style='List Bullet')

doc.add_paragraph('8.1.6 Facade Pattern')
doc.add_paragraph('Location: backend/app/services/rag_service.py (RAGService class)', style='List Bullet')
doc.add_paragraph('Implementation: RAGService provides a simplified interface to the complex RAG subsystem, coordinating PDFProcessor, EmbeddingModel, and LLMClient.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Hides complexity of multi-step RAG pipeline from clients', style='List Bullet')
doc.add_paragraph('  - Single entry point for document indexing and answer generation', style='List Bullet')
doc.add_paragraph('  - Reduces coupling between route handlers and utility classes', style='List Bullet')

doc.add_paragraph('8.1.7 Active Record Pattern')
doc.add_paragraph('Location: backend/app/models/ (User, Question, Document, DocumentChunk)', style='List Bullet')
doc.add_paragraph('Implementation: SQLAlchemy ORM models encapsulate both data and database operations.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Objects carry their own persistence logic', style='List Bullet')
doc.add_paragraph('  - Intuitive API for CRUD operations (model.query, db.session.add)', style='List Bullet')
doc.add_paragraph('  - Automatic SQL generation and parameterization', style='List Bullet')

doc.add_paragraph('8.1.8 Data Transfer Object (DTO) Pattern')
doc.add_paragraph('Location: backend/app/models/*.py (to_dict() methods)', style='List Bullet')
doc.add_paragraph('Implementation: Each model implements to_dict() to convert ORM objects to JSON-serializable dictionaries.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Clean serialization layer for API responses', style='List Bullet')
doc.add_paragraph('  - Control over which attributes are exposed to clients', style='List Bullet')
doc.add_paragraph('  - Handles type conversion (UUID to string, datetime to ISO format)', style='List Bullet')

# 8.2 Frontend Design Patterns
add_heading_with_bookmark(doc, '8.2 Frontend Design Patterns', level=2)

doc.add_paragraph('8.2.1 Provider Pattern')
doc.add_paragraph('Location: frontend/src/App.jsx (AuthContext.Provider)', style='List Bullet')
doc.add_paragraph('Implementation: React Context Provider wraps the application to provide global authentication state.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Avoids prop drilling through component hierarchy', style='List Bullet')
doc.add_paragraph('  - Centralized state management for authentication', style='List Bullet')
doc.add_paragraph('  - Any component can access auth state via useAuth() hook', style='List Bullet')

doc.add_paragraph('8.2.2 Custom Hook Pattern')
doc.add_paragraph('Location: frontend/src/App.jsx (useAuth hook)', style='List Bullet')
doc.add_paragraph('Implementation: Custom hook useAuth() provides access to AuthContext with error handling.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Encapsulates context access logic', style='List Bullet')
doc.add_paragraph('  - Provides meaningful error if used outside Provider', style='List Bullet')
doc.add_paragraph('  - Clean, reusable API for components', style='List Bullet')

doc.add_paragraph('8.2.3 Container/Presentational Pattern')
doc.add_paragraph('Location: frontend/src/pages/ (containers) vs frontend/src/components/ (presentational)', style='List Bullet')
doc.add_paragraph('Implementation: Page components (HomePage, LoginPage) manage state and logic; UI components (QuestionItem, Toast) focus on presentation.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Separation of concerns between logic and presentation', style='List Bullet')
doc.add_paragraph('  - Presentational components are easily reusable and testable', style='List Bullet')
doc.add_paragraph('  - Container components orchestrate data flow', style='List Bullet')

doc.add_paragraph('8.2.4 Composition Pattern')
doc.add_paragraph('Location: frontend/src/pages/HomePage.jsx', style='List Bullet')
doc.add_paragraph('Implementation: HomePage composes multiple child components (Header, QuestionList, QuestionForm, Modal, Toast).', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Complex UIs built from simple, focused components', style='List Bullet')
doc.add_paragraph('  - Each component has single responsibility', style='List Bullet')
doc.add_paragraph('  - Easy to modify or replace individual components', style='List Bullet')

doc.add_paragraph('8.2.5 Module Pattern')
doc.add_paragraph('Location: frontend/src/services/api.js', style='List Bullet')
doc.add_paragraph('Implementation: API clients (authApi, questionsApi, ragApi) exported as module objects with related functions.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Logical grouping of related API calls', style='List Bullet')
doc.add_paragraph('  - Encapsulation of HTTP request details', style='List Bullet')
doc.add_paragraph('  - Clean import syntax (import { questionsApi } from "./services/api")', style='List Bullet')

doc.add_paragraph('8.2.6 Observer Pattern')
doc.add_paragraph('Location: Throughout React components (useState, useEffect)', style='List Bullet')
doc.add_paragraph('Implementation: React state hooks implement observer pattern - state changes trigger component re-renders.', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Automatic UI updates when state changes', style='List Bullet')
doc.add_paragraph('  - Declarative data flow', style='List Bullet')
doc.add_paragraph('  - Side effects managed via useEffect subscriptions', style='List Bullet')

# 8.3 Architectural Patterns
add_heading_with_bookmark(doc, '8.3 Architectural Patterns', level=2)

doc.add_paragraph('8.3.1 Three-Tier Architecture')
doc.add_paragraph('Implementation: The application is organized into three distinct layers:', style='List Bullet')
doc.add_paragraph('  - Presentation Tier: React frontend handles user interface and interactions', style='List Bullet')
doc.add_paragraph('  - Business Logic Tier: Flask services process requests and implement business rules', style='List Bullet')
doc.add_paragraph('  - Data Tier: PostgreSQL database stores persistent data', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Independent scaling of each tier', style='List Bullet')
doc.add_paragraph('  - Technology flexibility (can replace any tier independently)', style='List Bullet')
doc.add_paragraph('  - Clear separation of responsibilities', style='List Bullet')

doc.add_paragraph('8.3.2 REST API Architecture')
doc.add_paragraph('Implementation: Backend exposes RESTful HTTP endpoints following standard conventions:', style='List Bullet')
doc.add_paragraph('  - GET for retrieval, POST for creation, PUT for update, DELETE for removal', style='List Bullet')
doc.add_paragraph('  - Resource-based URLs (/api/questions, /api/questions/{id})', style='List Bullet')
doc.add_paragraph('  - Stateless requests with authentication via headers', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Standard, predictable API design', style='List Bullet')
doc.add_paragraph('  - Easy integration with any HTTP client', style='List Bullet')
doc.add_paragraph('  - Cacheable responses, scalable architecture', style='List Bullet')

doc.add_paragraph('8.3.3 MVC-like Pattern')
doc.add_paragraph('Implementation: While not strict MVC, the application follows similar separation:', style='List Bullet')
doc.add_paragraph('  - Model: SQLAlchemy ORM classes and service layer', style='List Bullet')
doc.add_paragraph('  - View: React components rendering the UI', style='List Bullet')
doc.add_paragraph('  - Controller: Flask route handlers coordinating requests/responses', style='List Bullet')
doc.add_paragraph('Benefits:', style='List Bullet')
doc.add_paragraph('  - Familiar pattern for developers', style='List Bullet')
doc.add_paragraph('  - Clear data flow through the application', style='List Bullet')
doc.add_paragraph('  - Testable components at each layer', style='List Bullet')

# 8.4 Pattern Summary Table
add_heading_with_bookmark(doc, '8.4 Pattern Summary', level=2)

doc.add_paragraph(
    'The following table summarizes all design patterns used in the Ask Holmes application:'
)

# Create a table for pattern summary
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'

# Header row
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Pattern'
hdr_cells[1].text = 'Category'
hdr_cells[2].text = 'Layer'
hdr_cells[3].text = 'Primary Benefit'

# Make header bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

# Data rows
patterns = [
    ('Application Factory', 'Creational', 'Backend', 'Configurable app instances'),
    ('Singleton', 'Creational', 'Backend', 'Memory-efficient ML model'),
    ('Service Layer', 'Architectural', 'Backend', 'Separation of concerns'),
    ('Repository', 'Structural', 'Backend', 'Data access abstraction'),
    ('Blueprint', 'Structural', 'Backend', 'Modular route organization'),
    ('Facade', 'Structural', 'Backend', 'Simplified RAG interface'),
    ('Active Record', 'Architectural', 'Backend', 'ORM data persistence'),
    ('DTO', 'Structural', 'Backend', 'Clean API serialization'),
    ('Provider', 'Behavioral', 'Frontend', 'Global state management'),
    ('Custom Hook', 'Behavioral', 'Frontend', 'Reusable stateful logic'),
    ('Container/Presentational', 'Structural', 'Frontend', 'UI/logic separation'),
    ('Composition', 'Structural', 'Frontend', 'Modular UI building'),
    ('Module', 'Structural', 'Frontend', 'Organized API clients'),
    ('Observer', 'Behavioral', 'Frontend', 'Reactive UI updates'),
    ('Three-Tier', 'Architectural', 'Full Stack', 'Scalable architecture'),
    ('REST API', 'Architectural', 'Full Stack', 'Standard HTTP interface'),
]

for pattern, category, layer, benefit in patterns:
    row_cells = table.add_row().cells
    row_cells[0].text = pattern
    row_cells[1].text = category
    row_cells[2].text = layer
    row_cells[3].text = benefit

# ============================================
# 9. UNIT TEST CASES
# ============================================
add_heading_with_bookmark(doc, '9. Unit Test Cases', level=1)

doc.add_paragraph(
    'This chapter documents the unit test cases for the Ask Holmes application. '
    'Unit tests verify that individual components function correctly in isolation. '
    'The test cases are organized by application layer and component, covering '
    'models, services, utilities, API routes, and frontend components.'
)

# 9.1 Backend Model Tests
add_heading_with_bookmark(doc, '9.1 Backend Model Tests', level=2)

doc.add_paragraph('9.1.1 User Model Tests')

user_table = doc.add_table(rows=1, cols=4)
user_table.style = 'Table Grid'
hdr = user_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

user_tests = [
    ('UM-001', 'Create user with valid email', 'email="test@example.com"', 'User created with UUID, email stored, created_at set'),
    ('UM-002', 'Create user with duplicate email', 'email="existing@example.com"', 'IntegrityError raised (unique constraint)'),
    ('UM-003', 'Create user with empty email', 'email=""', 'Validation error or database constraint failure'),
    ('UM-004', 'User to_dict() serialization', 'Valid User object', 'Dict with id (string), email, created_at (ISO format)'),
    ('UM-005', 'User-Question relationship', 'User with 3 questions', 'user.questions returns list of 3 Question objects'),
    ('UM-006', 'Cascade delete user', 'Delete user with questions', 'User and all associated questions deleted'),
]

for test_id, case, input_val, expected in user_tests:
    row = user_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.1.2 Question Model Tests')

question_table = doc.add_table(rows=1, cols=4)
question_table.style = 'Table Grid'
hdr = question_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

question_tests = [
    ('QM-001', 'Create question with valid data', 'user_id, question="What is...?"', 'Question created with UUID, timestamps set'),
    ('QM-002', 'Create question without user_id', 'question="Test?" (no user_id)', 'IntegrityError (foreign key constraint)'),
    ('QM-003', 'Create question with answer', 'question, answer="The answer..."', 'Question created with answer field populated'),
    ('QM-004', 'Create question without answer', 'question only', 'Question created with answer=None'),
    ('QM-005', 'Update question updates timestamp', 'Modify existing question', 'updated_at timestamp changes, created_at unchanged'),
    ('QM-006', 'Question to_dict() serialization', 'Valid Question object', 'Dict with all fields, UUIDs as strings, dates as ISO'),
    ('QM-007', 'Question-User backref', 'Question with user_id', 'question.user returns associated User object'),
]

for test_id, case, input_val, expected in question_tests:
    row = question_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.1.3 Document Model Tests')

document_table = doc.add_table(rows=1, cols=4)
document_table.style = 'Table Grid'
hdr = document_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

document_tests = [
    ('DM-001', 'Create document with valid data', 'filename, title, file_hash', 'Document created with UUID, indexed_at set'),
    ('DM-002', 'Create document with duplicate hash', 'Existing file_hash', 'IntegrityError (unique constraint on hash)'),
    ('DM-003', 'Document chunk_count default', 'Create without chunk_count', 'chunk_count defaults to 0'),
    ('DM-004', 'Document to_dict() serialization', 'Valid Document object', 'Dict with all fields serialized correctly'),
    ('DM-005', 'Document-Chunk relationship', 'Document with 5 chunks', 'document.chunks returns list of 5 DocumentChunk objects'),
    ('DM-006', 'Cascade delete document', 'Delete document with chunks', 'Document and all associated chunks deleted'),
]

for test_id, case, input_val, expected in document_tests:
    row = document_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.1.4 DocumentChunk Model Tests')

chunk_table = doc.add_table(rows=1, cols=4)
chunk_table.style = 'Table Grid'
hdr = chunk_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

chunk_tests = [
    ('CM-001', 'Create chunk with valid data', 'document_id, chunk_text, chunk_index', 'Chunk created with UUID, created_at set'),
    ('CM-002', 'Create chunk with embedding', 'chunk_text, 384-dim vector', 'Chunk created with embedding stored'),
    ('CM-003', 'Create chunk without document_id', 'chunk_text only', 'IntegrityError (foreign key constraint)'),
    ('CM-004', 'Chunk to_dict() serialization', 'Valid DocumentChunk', 'Dict with fields (embedding excluded)'),
    ('CM-005', 'Chunk-Document backref', 'Chunk with document_id', 'chunk.document returns associated Document'),
    ('CM-006', 'Vector similarity query', 'Query embedding vector', 'Chunks ordered by cosine distance'),
]

for test_id, case, input_val, expected in chunk_tests:
    row = chunk_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

# 9.2 Backend Service Tests
add_heading_with_bookmark(doc, '9.2 Backend Service Tests', level=2)

doc.add_paragraph('9.2.1 AuthService Tests')

auth_table = doc.add_table(rows=1, cols=4)
auth_table.style = 'Table Grid'
hdr = auth_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

auth_tests = [
    ('AS-001', 'Authenticate with valid email', 'Registered user email', '(User object, None) returned'),
    ('AS-002', 'Authenticate with unregistered email', 'Unknown email', '(None, "User not registered...") returned'),
    ('AS-003', 'Authenticate with empty email', 'email=""', '(None, "Email is required") returned'),
    ('AS-004', 'Authenticate with invalid format', 'email="invalid"', '(None, "Invalid email format") returned'),
    ('AS-005', 'Authenticate normalizes email', 'email="TEST@Example.COM"', 'Email converted to lowercase before lookup'),
    ('AS-006', 'Get user by valid ID', 'Existing user UUID', 'User object returned'),
    ('AS-007', 'Get user by invalid ID', 'Non-existent UUID', 'None returned'),
    ('AS-008', 'Create user with valid email', 'New unique email', '(User object, None) returned'),
    ('AS-009', 'Create user with existing email', 'Already registered email', '(None, "User already exists") returned'),
    ('AS-010', 'Create user normalizes email', 'email="NEW@Example.COM"', 'Email stored as lowercase'),
]

for test_id, case, input_val, expected in auth_tests:
    row = auth_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.2.2 QuestionService Tests')

qs_table = doc.add_table(rows=1, cols=4)
qs_table.style = 'Table Grid'
hdr = qs_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

qs_tests = [
    ('QS-001', 'Get questions for user with questions', 'user_id with 5 questions', 'List of 5 questions, ordered by created_at desc'),
    ('QS-002', 'Get questions for user with none', 'user_id with 0 questions', 'Empty list returned'),
    ('QS-003', 'Get question by ID (owned)', 'question_id, matching user_id', 'Question object returned'),
    ('QS-004', 'Get question by ID (not owned)', 'question_id, different user_id', 'None returned (ownership check)'),
    ('QS-005', 'Get question by invalid ID', 'Non-existent question_id', 'None returned'),
    ('QS-006', 'Create question with text only', 'user_id, question text', '(Question, None) with answer=None'),
    ('QS-007', 'Create question with answer', 'user_id, question, answer', '(Question, None) with answer populated'),
    ('QS-008', 'Create question with empty text', 'user_id, question=""', '(None, "Question text is required")'),
    ('QS-009', 'Update question text', 'question_id, new question text', '(Updated Question, None), updated_at changed'),
    ('QS-010', 'Update question answer', 'question_id, new answer', '(Updated Question, None), answer updated'),
    ('QS-011', 'Update non-existent question', 'Invalid question_id', '(None, "Question not found")'),
    ('QS-012', 'Update question not owned', 'question_id, wrong user_id', '(None, "Question not found")'),
    ('QS-013', 'Delete existing question', 'Valid question_id, user_id', '(True, None) question removed'),
    ('QS-014', 'Delete non-existent question', 'Invalid question_id', '(False, "Question not found")'),
    ('QS-015', 'Delete question not owned', 'question_id, wrong user_id', '(False, "Question not found")'),
]

for test_id, case, input_val, expected in qs_tests:
    row = qs_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.2.3 RAGService Tests')

rag_table = doc.add_table(rows=1, cols=4)
rag_table.style = 'Table Grid'
hdr = rag_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rag_tests = [
    ('RS-001', 'Index valid PDF document', 'Path to valid PDF file', '(Document, None) with chunks created'),
    ('RS-002', 'Index non-existent PDF', 'Invalid file path', '(None, "PDF file not found...")'),
    ('RS-003', 'Index duplicate document', 'Already indexed PDF (same hash)', '(None, "Document already indexed...")'),
    ('RS-004', 'Index all documents (with PDFs)', 'Books folder with 3 PDFs', '(3 successes, 0 failures)'),
    ('RS-005', 'Index all documents (empty folder)', 'Empty books folder', '(0 successes, 1 failure message)'),
    ('RS-006', 'Search similar chunks', 'Query string, top_k=3', 'List of 3 most similar DocumentChunks'),
    ('RS-007', 'Search with no indexed docs', 'Query string, empty database', 'Empty list returned'),
    ('RS-008', 'Generate answer (docs exist)', 'Question with relevant docs', '(Answer string, None)'),
    ('RS-009', 'Generate answer (no docs)', 'Question, empty database', '(None, "No relevant documents found...")'),
    ('RS-010', 'Get indexed documents', 'Database with 5 documents', 'List of 5 Documents, ordered by indexed_at desc'),
    ('RS-011', 'Delete existing document', 'Valid document_id', '(True, None), document and chunks removed'),
    ('RS-012', 'Delete non-existent document', 'Invalid document_id', '(False, "Document not found")'),
]

for test_id, case, input_val, expected in rag_tests:
    row = rag_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

# 9.3 Backend Utility Tests
add_heading_with_bookmark(doc, '9.3 Backend Utility Tests', level=2)

doc.add_paragraph('9.3.1 EmbeddingModel Tests')

emb_table = doc.add_table(rows=1, cols=4)
emb_table.style = 'Table Grid'
hdr = emb_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

emb_tests = [
    ('EM-001', 'Singleton pattern enforced', 'Create two instances', 'Both variables reference same instance'),
    ('EM-002', 'Embed single text', '"Hello world"', 'List of 384 floats returned'),
    ('EM-003', 'Embed empty string', '""', 'List of 384 floats (zero or near-zero values)'),
    ('EM-004', 'Embed batch of texts', '["text1", "text2", "text3"]', 'List of 3 embeddings, each 384 floats'),
    ('EM-005', 'Embed batch empty list', '[]', 'Empty list returned'),
    ('EM-006', 'Embedding consistency', 'Same text twice', 'Identical embeddings returned'),
]

for test_id, case, input_val, expected in emb_tests:
    row = emb_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.3.2 PDFProcessor Tests')

pdf_table = doc.add_table(rows=1, cols=4)
pdf_table.style = 'Table Grid'
hdr = pdf_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

pdf_tests = [
    ('PP-001', 'Extract text from valid PDF', 'Path to text-based PDF', 'String containing PDF text content'),
    ('PP-002', 'Extract text from empty PDF', 'PDF with no text', 'Empty string returned'),
    ('PP-003', 'Clean text removes extra whitespace', '"Hello   world\\n\\n"', '"Hello world"'),
    ('PP-004', 'Chunk text shorter than chunk_size', '100 char text, chunk_size=500', 'List with 1 chunk'),
    ('PP-005', 'Chunk text creates overlapping chunks', '1500 char text', 'Multiple chunks with overlap at boundaries'),
    ('PP-006', 'Chunk text respects sentence boundaries', 'Text with sentences', 'Chunks end at periods when possible'),
    ('PP-007', 'Calculate file hash', 'Valid PDF file', '64-character SHA-256 hex string'),
    ('PP-008', 'Calculate hash consistency', 'Same file twice', 'Identical hash returned'),
    ('PP-009', 'Process PDF returns metadata', 'Valid PDF file', 'Dict with filename, title, file_hash, chunks, chunk_count'),
    ('PP-010', 'Process non-existent PDF', 'Invalid path', 'FileNotFoundError raised'),
]

for test_id, case, input_val, expected in pdf_tests:
    row = pdf_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.3.3 LLMClient Tests')

llm_table = doc.add_table(rows=1, cols=4)
llm_table.style = 'Table Grid'
hdr = llm_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Input'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

llm_tests = [
    ('LC-001', 'Generate answer with context', 'Question, list of context chunks', 'Non-empty answer string'),
    ('LC-002', 'Generate answer empty context', 'Question, empty context list', 'Answer indicating no context available'),
    ('LC-003', 'API error handling', 'Invalid API key', 'Appropriate exception raised'),
    ('LC-004', 'Context formatting', 'Multiple context chunks', 'Chunks joined with separator in prompt'),
    ('LC-005', 'System prompt included', 'Any question', 'Sherlock Holmes expert prompt used'),
]

for test_id, case, input_val, expected in llm_tests:
    row = llm_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = input_val
    row[3].text = expected

# 9.4 Backend API Route Tests
add_heading_with_bookmark(doc, '9.4 Backend API Route Tests', level=2)

doc.add_paragraph('9.4.1 Authentication Routes (/api/auth)')

auth_route_table = doc.add_table(rows=1, cols=5)
auth_route_table.style = 'Table Grid'
hdr = auth_route_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Endpoint'
hdr[2].text = 'Test Case'
hdr[3].text = 'Input'
hdr[4].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

auth_route_tests = [
    ('AR-001', 'POST /login', 'Login with valid email', '{"email": "user@test.com"}', '200 OK, user object returned'),
    ('AR-002', 'POST /login', 'Login with unregistered email', '{"email": "unknown@test.com"}', '401 Unauthorized, error message'),
    ('AR-003', 'POST /login', 'Login with missing email', '{}', '400 Bad Request'),
    ('AR-004', 'POST /login', 'Login with invalid JSON', 'Invalid JSON body', '400 Bad Request'),
    ('AR-005', 'POST /register', 'Register new user', '{"email": "new@test.com"}', '201 Created, user object'),
    ('AR-006', 'POST /register', 'Register existing email', '{"email": "existing@test.com"}', '400 Bad Request, error'),
    ('AR-007', 'GET /user/<id>', 'Get existing user', 'Valid user UUID', '200 OK, user object'),
    ('AR-008', 'GET /user/<id>', 'Get non-existent user', 'Invalid UUID', '404 Not Found'),
]

for test_id, endpoint, case, input_val, expected in auth_route_tests:
    row = auth_route_table.add_row().cells
    row[0].text = test_id
    row[1].text = endpoint
    row[2].text = case
    row[3].text = input_val
    row[4].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.4.2 Questions Routes (/api/questions)')

q_route_table = doc.add_table(rows=1, cols=5)
q_route_table.style = 'Table Grid'
hdr = q_route_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Endpoint'
hdr[2].text = 'Test Case'
hdr[3].text = 'Input'
hdr[4].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

q_route_tests = [
    ('QR-001', 'GET /', 'List questions with header', 'X-User-ID header set', '200 OK, questions array'),
    ('QR-002', 'GET /', 'List questions no header', 'No X-User-ID header', '401 Unauthorized'),
    ('QR-003', 'GET /<id>', 'Get owned question', 'Valid question ID, owner', '200 OK, question object'),
    ('QR-004', 'GET /<id>', 'Get unowned question', 'Valid ID, different user', '404 Not Found'),
    ('QR-005', 'POST /', 'Create question', '{"question": "What is...?"}', '201 Created, question object'),
    ('QR-006', 'POST /', 'Create with answer', '{"question": "...", "answer": "..."}', '201 Created'),
    ('QR-007', 'POST /', 'Create empty question', '{"question": ""}', '400 Bad Request'),
    ('QR-008', 'POST /', 'Create no body', 'Empty request body', '400 Bad Request'),
    ('QR-009', 'PUT /<id>', 'Update question text', '{"question": "Updated?"}', '200 OK, updated question'),
    ('QR-010', 'PUT /<id>', 'Update answer', '{"answer": "New answer"}', '200 OK, updated question'),
    ('QR-011', 'PUT /<id>', 'Update non-existent', 'Invalid question ID', '400 Bad Request'),
    ('QR-012', 'DELETE /<id>', 'Delete owned question', 'Valid question ID, owner', '200 OK, success message'),
    ('QR-013', 'DELETE /<id>', 'Delete unowned question', 'Valid ID, different user', '400 Bad Request'),
]

for test_id, endpoint, case, input_val, expected in q_route_tests:
    row = q_route_table.add_row().cells
    row[0].text = test_id
    row[1].text = endpoint
    row[2].text = case
    row[3].text = input_val
    row[4].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.4.3 RAG Routes (/api/rag)')

rag_route_table = doc.add_table(rows=1, cols=5)
rag_route_table.style = 'Table Grid'
hdr = rag_route_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Endpoint'
hdr[2].text = 'Test Case'
hdr[3].text = 'Input'
hdr[4].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

rag_route_tests = [
    ('RR-001', 'POST /index', 'Index documents success', 'Books folder with PDFs', '200 OK, successes/failures'),
    ('RR-002', 'POST /index', 'Index empty folder', 'Empty books folder', '200 OK, failure message'),
    ('RR-003', 'POST /query', 'Query with question', '{"question": "Who is Holmes?"}', '200 OK, answer string'),
    ('RR-004', 'POST /query', 'Query empty question', '{"question": ""}', '400 Bad Request'),
    ('RR-005', 'POST /query', 'Query no documents', 'Empty database', '400 Bad Request, no docs message'),
    ('RR-006', 'POST /search', 'Search chunks', '{"query": "detective", "top_k": 5}', '200 OK, chunk array'),
    ('RR-007', 'GET /documents', 'List indexed documents', 'Database with documents', '200 OK, documents array'),
    ('RR-008', 'GET /documents', 'List empty', 'Empty database', '200 OK, empty array'),
    ('RR-009', 'DELETE /docs/<id>', 'Delete document', 'Valid document ID', '200 OK, success message'),
    ('RR-010', 'DELETE /docs/<id>', 'Delete non-existent', 'Invalid document ID', '400 Bad Request'),
]

for test_id, endpoint, case, input_val, expected in rag_route_tests:
    row = rag_route_table.add_row().cells
    row[0].text = test_id
    row[1].text = endpoint
    row[2].text = case
    row[3].text = input_val
    row[4].text = expected

# 9.5 Frontend Component Tests
add_heading_with_bookmark(doc, '9.5 Frontend Component Tests', level=2)

doc.add_paragraph('9.5.1 App Component Tests')

app_table = doc.add_table(rows=1, cols=4)
app_table.style = 'Table Grid'
hdr = app_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Condition'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

app_tests = [
    ('FC-001', 'Renders login when not authenticated', 'user = null', 'LoginPage component rendered'),
    ('FC-002', 'Renders home when authenticated', 'user = {id, email}', 'HomePage component rendered'),
    ('FC-003', 'Shows loading spinner initially', 'loading = true', 'Spinner element displayed'),
    ('FC-004', 'Loads user from localStorage', 'Valid user in localStorage', 'User state populated from storage'),
    ('FC-005', 'Handles invalid localStorage data', 'Invalid JSON in localStorage', 'localStorage cleared, user = null'),
    ('FC-006', 'Login function updates state', 'Call login(userData)', 'User state set, localStorage updated'),
    ('FC-007', 'Logout function clears state', 'Call logout()', 'User = null, localStorage cleared'),
    ('FC-008', 'AuthContext provides value', 'useAuth() in child', 'Returns {user, login, logout}'),
]

for test_id, case, condition, expected in app_tests:
    row = app_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = condition
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.5.2 LoginPage Tests')

login_table = doc.add_table(rows=1, cols=4)
login_table.style = 'Table Grid'
hdr = login_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Action'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

login_tests = [
    ('LP-001', 'Renders email input', 'Component mounts', 'Email input field visible'),
    ('LP-002', 'Renders submit button', 'Component mounts', 'Submit button visible'),
    ('LP-003', 'Email input updates state', 'Type in email field', 'email state matches input value'),
    ('LP-004', 'Submit disabled when empty', 'email = ""', 'Submit button disabled'),
    ('LP-005', 'Submit enabled with email', 'email = "test@test.com"', 'Submit button enabled'),
    ('LP-006', 'Shows loading on submit', 'Form submitted', 'Loading spinner shown, button disabled'),
    ('LP-007', 'Calls login on success', 'API returns user', 'login() called with user data'),
    ('LP-008', 'Shows error on failure', 'API returns error', 'Error message displayed'),
    ('LP-009', 'Clears error on new submit', 'Submit after error', 'Error message cleared'),
]

for test_id, case, action, expected in login_tests:
    row = login_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = action
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.5.3 HomePage Tests')

home_table = doc.add_table(rows=1, cols=4)
home_table.style = 'Table Grid'
hdr = home_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Condition'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

home_tests = [
    ('HP-001', 'Fetches questions on mount', 'Component mounts', 'questionsApi.getAll() called'),
    ('HP-002', 'Displays loading state', 'loading = true', 'Loading spinner displayed'),
    ('HP-003', 'Displays error state', 'API fetch fails', 'Error message and retry button'),
    ('HP-004', 'Renders question list', 'Questions loaded', 'QuestionList with questions prop'),
    ('HP-005', 'Opens modal on add click', 'Click "Add Question"', 'Modal opens with QuestionForm'),
    ('HP-006', 'Creates question successfully', 'Submit new question', 'Question added to list, toast shown'),
    ('HP-007', 'Opens edit modal', 'Click edit on question', 'Modal opens with question data'),
    ('HP-008', 'Updates question successfully', 'Submit edited question', 'Question updated in list, toast shown'),
    ('HP-009', 'Deletes question', 'Confirm delete', 'Question removed from list, toast shown'),
    ('HP-010', 'Ask documents returns answer', 'Click Ask Documents', 'Answer populated in form'),
]

for test_id, case, condition, expected in home_tests:
    row = home_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = condition
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.5.4 QuestionForm Tests')

form_table = doc.add_table(rows=1, cols=4)
form_table.style = 'Table Grid'
hdr = form_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Test Case'
hdr[2].text = 'Action'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

form_tests = [
    ('QF-001', 'Renders empty form for create', 'editingQuestion = null', 'Empty question and answer fields'),
    ('QF-002', 'Populates form for edit', 'editingQuestion = {...}', 'Fields populated with question data'),
    ('QF-003', 'Submit disabled when empty', 'question = ""', 'Submit button disabled'),
    ('QF-004', 'Submit enabled with question', 'question = "Test?"', 'Submit button enabled'),
    ('QF-005', 'Calls onSubmit for create', 'Submit new question', 'onSubmit(question, answer) called'),
    ('QF-006', 'Calls onSubmit for edit', 'Submit edited question', 'onSubmit(id, question, answer) called'),
    ('QF-007', 'Ask Documents button works', 'Click Ask Documents', 'onAskDocuments called, answer populated'),
    ('QF-008', 'Shows loading during ask', 'Asking documents', 'Spinner shown, button disabled'),
    ('QF-009', 'Cancel clears form', 'Click cancel', 'onCancelEdit called, form reset'),
]

for test_id, case, action, expected in form_tests:
    row = form_table.add_row().cells
    row[0].text = test_id
    row[1].text = case
    row[2].text = action
    row[3].text = expected

doc.add_paragraph('')
doc.add_paragraph('9.5.5 Common Component Tests')

common_table = doc.add_table(rows=1, cols=4)
common_table.style = 'Table Grid'
hdr = common_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Component'
hdr[2].text = 'Test Case'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

common_tests = [
    ('CC-001', 'Header', 'Displays user email', 'User email shown in header'),
    ('CC-002', 'Header', 'Logout button works', 'logout() called on click'),
    ('CC-003', 'Modal', 'Not rendered when closed', 'isOpen=false returns null'),
    ('CC-004', 'Modal', 'Rendered when open', 'isOpen=true shows overlay and content'),
    ('CC-005', 'Modal', 'Closes on backdrop click', 'onClose called when clicking overlay'),
    ('CC-006', 'Modal', 'Closes on Escape key', 'onClose called on Escape keypress'),
    ('CC-007', 'Toast', 'Displays message', 'Message text visible'),
    ('CC-008', 'Toast', 'Shows success style', 'type="success" applies success class'),
    ('CC-009', 'Toast', 'Shows error style', 'type="error" applies error class'),
    ('CC-010', 'QuestionList', 'Renders empty state', 'questions=[] shows empty message'),
    ('CC-011', 'QuestionList', 'Renders questions', 'questions=[...] renders QuestionItems'),
    ('CC-012', 'QuestionItem', 'Displays question text', 'Question content visible'),
    ('CC-013', 'QuestionItem', 'Shows truncated answer', 'Long answer truncated with ellipsis'),
    ('CC-014', 'QuestionItem', 'Two-step delete', 'First click shows confirm, second deletes'),
]

for test_id, component, case, expected in common_tests:
    row = common_table.add_row().cells
    row[0].text = test_id
    row[1].text = component
    row[2].text = case
    row[3].text = expected

# 9.6 Frontend API Service Tests
add_heading_with_bookmark(doc, '9.6 Frontend API Service Tests', level=2)

api_table = doc.add_table(rows=1, cols=4)
api_table.style = 'Table Grid'
hdr = api_table.rows[0].cells
hdr[0].text = 'Test ID'
hdr[1].text = 'Service Method'
hdr[2].text = 'Test Case'
hdr[3].text = 'Expected Result'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

api_tests = [
    ('API-001', 'authApi.login', 'Successful login', 'Returns user object from response'),
    ('API-002', 'authApi.login', 'Failed login', 'Throws ApiError with message'),
    ('API-003', 'questionsApi.getAll', 'Fetch questions', 'Returns questions array'),
    ('API-004', 'questionsApi.getAll', 'Includes X-User-ID header', 'Header sent with request'),
    ('API-005', 'questionsApi.create', 'Create question', 'Returns created question'),
    ('API-006', 'questionsApi.update', 'Update question', 'Returns updated question'),
    ('API-007', 'questionsApi.delete', 'Delete question', 'Returns success response'),
    ('API-008', 'ragApi.query', 'Query documents', 'Returns answer string'),
    ('API-009', 'ragApi.indexDocuments', 'Index documents', 'Returns success/failure lists'),
    ('API-010', 'ApiError', 'Network error', 'ApiError with status 0, "Network error" message'),
    ('API-011', 'ApiError', 'Server error', 'ApiError with status code and error message'),
]

for test_id, method, case, expected in api_tests:
    row = api_table.add_row().cells
    row[0].text = test_id
    row[1].text = method
    row[2].text = case
    row[3].text = expected

# 9.7 Test Summary
add_heading_with_bookmark(doc, '9.7 Test Summary', level=2)

doc.add_paragraph(
    'The following table summarizes the total number of test cases by category:'
)

summary_table = doc.add_table(rows=1, cols=3)
summary_table.style = 'Table Grid'
hdr = summary_table.rows[0].cells
hdr[0].text = 'Category'
hdr[1].text = 'Subcategory'
hdr[2].text = 'Test Count'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

summary_data = [
    ('Backend Models', 'User, Question, Document, DocumentChunk', '25'),
    ('Backend Services', 'AuthService, QuestionService, RAGService', '37'),
    ('Backend Utilities', 'EmbeddingModel, PDFProcessor, LLMClient', '21'),
    ('Backend API Routes', 'Auth, Questions, RAG endpoints', '31'),
    ('Frontend Components', 'App, LoginPage, HomePage, QuestionForm, Common', '50'),
    ('Frontend API Services', 'authApi, questionsApi, ragApi', '11'),
    ('Total', '', '175'),
]

for category, subcat, count in summary_data:
    row = summary_table.add_row().cells
    row[0].text = category
    row[1].text = subcat
    row[2].text = count

# ============================================
# 10. ARCHITECTURAL DESIGN
# ============================================
add_heading_with_bookmark(doc, '10. Architectural Design', level=1)

doc.add_paragraph(
    'This chapter presents the architectural design of the Ask Holmes application, '
    'describing the system structure, component organization, and interactions between '
    'different layers. The architecture follows industry best practices including '
    'separation of concerns, modularity, and scalability.'
)

# 10.1 Architecture Overview
add_heading_with_bookmark(doc, '10.1 Architecture Overview', level=2)

doc.add_paragraph(
    'The Ask Holmes application implements a modern three-tier web architecture with '
    'clear separation between the presentation, business logic, and data layers. '
    'The system is designed as a client-server application where the React frontend '
    'communicates with the Flask backend through RESTful APIs.'
)

doc.add_paragraph('Key Architectural Characteristics:')
doc.add_paragraph('Layered Architecture: Clear separation into Client, Server, and Data layers', style='List Bullet')
doc.add_paragraph('Microservice-Ready: Modular design allows easy extraction into microservices', style='List Bullet')
doc.add_paragraph('Stateless Backend: RESTful API design enables horizontal scaling', style='List Bullet')
doc.add_paragraph('Event-Driven Frontend: React components respond to state changes reactively', style='List Bullet')
doc.add_paragraph('External Service Integration: Loosely coupled integration with AI services', style='List Bullet')

# 10.2 Client Layer
add_heading_with_bookmark(doc, '10.2 Client Layer', level=2)

doc.add_paragraph(
    'The client layer consists of the web browser and the React single-page application (SPA). '
    'This layer is responsible for rendering the user interface, handling user interactions, '
    'and managing client-side state.'
)

doc.add_paragraph('10.2.1 Web Browser')
doc.add_paragraph('Entry point for end users accessing the application', style='List Bullet')
doc.add_paragraph('Renders the React application and executes JavaScript', style='List Bullet')
doc.add_paragraph('Stores authentication state in localStorage for session persistence', style='List Bullet')

doc.add_paragraph('10.2.2 React Frontend Components')

doc.add_paragraph('Pages:')
doc.add_paragraph('LoginPage: Handles user authentication flow', style='List Bullet')
doc.add_paragraph('HomePage: Main application interface for question management', style='List Bullet')

doc.add_paragraph('Components:')
doc.add_paragraph('Header: Application branding and user controls', style='List Bullet')
doc.add_paragraph('Modal: Reusable dialog container for forms', style='List Bullet')
doc.add_paragraph('Toast: Notification system for user feedback', style='List Bullet')
doc.add_paragraph('QuestionList/QuestionItem: Question display and management', style='List Bullet')
doc.add_paragraph('QuestionForm: Input form for creating/editing questions', style='List Bullet')

doc.add_paragraph('AuthContext:')
doc.add_paragraph('Global state management using React Context API', style='List Bullet')
doc.add_paragraph('Provides user authentication state to all components', style='List Bullet')
doc.add_paragraph('Manages login/logout operations and localStorage sync', style='List Bullet')

doc.add_paragraph('API Services:')
doc.add_paragraph('authApi: Authentication-related HTTP requests', style='List Bullet')
doc.add_paragraph('questionsApi: Question CRUD operations', style='List Bullet')
doc.add_paragraph('ragApi: RAG pipeline operations (indexing, querying)', style='List Bullet')

# 10.3 Server Layer
add_heading_with_bookmark(doc, '10.3 Server Layer', level=2)

doc.add_paragraph(
    'The server layer is built with Flask and follows a layered architecture pattern. '
    'It processes client requests, executes business logic, and manages data persistence.'
)

doc.add_paragraph('10.3.1 API Layer (Blueprints)')
doc.add_paragraph('Purpose: HTTP request handling, routing, and response formatting', style='List Bullet')
doc.add_paragraph('Implementation: Flask Blueprints for modular route organization', style='List Bullet')
doc.add_paragraph('Components:', style='List Bullet')
doc.add_paragraph('  - Auth Routes (/api/auth/*): Login, registration, user retrieval', style='List Bullet')
doc.add_paragraph('  - Question Routes (/api/questions/*): CRUD operations for questions', style='List Bullet')
doc.add_paragraph('  - RAG Routes (/api/rag/*): Document indexing and answer generation', style='List Bullet')
doc.add_paragraph('Responsibilities:', style='List Bullet')
doc.add_paragraph('  - Request validation and parsing', style='List Bullet')
doc.add_paragraph('  - Authentication header extraction (X-User-ID)', style='List Bullet')
doc.add_paragraph('  - Response formatting (JSON)', style='List Bullet')
doc.add_paragraph('  - HTTP status code management', style='List Bullet')

doc.add_paragraph('10.3.2 Service Layer')
doc.add_paragraph('Purpose: Business logic encapsulation and orchestration', style='List Bullet')
doc.add_paragraph('Implementation: Python classes with static methods', style='List Bullet')
doc.add_paragraph('Components:', style='List Bullet')
doc.add_paragraph('  - AuthService: User authentication and registration logic', style='List Bullet')
doc.add_paragraph('  - QuestionService: Question management business rules', style='List Bullet')
doc.add_paragraph('  - RAGService: RAG pipeline orchestration (Facade pattern)', style='List Bullet')
doc.add_paragraph('Responsibilities:', style='List Bullet')
doc.add_paragraph('  - Business rule enforcement', style='List Bullet')
doc.add_paragraph('  - Data validation beyond basic input checks', style='List Bullet')
doc.add_paragraph('  - Coordination between utilities and models', style='List Bullet')
doc.add_paragraph('  - Transaction management', style='List Bullet')

doc.add_paragraph('10.3.3 Utility Layer')
doc.add_paragraph('Purpose: Specialized functionality and external integrations', style='List Bullet')
doc.add_paragraph('Components:', style='List Bullet')
doc.add_paragraph('  - PDFProcessor: PDF text extraction and chunking', style='List Bullet')
doc.add_paragraph('  - EmbeddingModel (Singleton): Vector embedding generation', style='List Bullet')
doc.add_paragraph('  - LLMClient: Anthropic Claude API integration', style='List Bullet')
doc.add_paragraph('  - Config: Centralized configuration management', style='List Bullet')
doc.add_paragraph('Design Decisions:', style='List Bullet')
doc.add_paragraph('  - Singleton pattern for EmbeddingModel to conserve memory', style='List Bullet')
doc.add_paragraph('  - Environment-based configuration for deployment flexibility', style='List Bullet')
doc.add_paragraph('  - Abstracted external service calls for testability', style='List Bullet')

doc.add_paragraph('10.3.4 Model Layer (ORM)')
doc.add_paragraph('Purpose: Data representation and persistence abstraction', style='List Bullet')
doc.add_paragraph('Implementation: SQLAlchemy ORM with Flask-SQLAlchemy', style='List Bullet')
doc.add_paragraph('Components:', style='List Bullet')
doc.add_paragraph('  - User Model: User account data', style='List Bullet')
doc.add_paragraph('  - Question Model: Question and answer storage', style='List Bullet')
doc.add_paragraph('  - Document Model: Indexed PDF metadata', style='List Bullet')
doc.add_paragraph('  - DocumentChunk Model: Vectorized text segments', style='List Bullet')
doc.add_paragraph('Features:', style='List Bullet')
doc.add_paragraph('  - Automatic schema generation', style='List Bullet')
doc.add_paragraph('  - Relationship mapping (one-to-many)', style='List Bullet')
doc.add_paragraph('  - Cascade delete for referential integrity', style='List Bullet')
doc.add_paragraph('  - JSON serialization via to_dict() methods', style='List Bullet')

# 10.4 Data Layer
add_heading_with_bookmark(doc, '10.4 Data Layer', level=2)

doc.add_paragraph(
    'The data layer uses PostgreSQL as the relational database management system, '
    'extended with pgvector for vector similarity search capabilities.'
)

doc.add_paragraph('10.4.1 PostgreSQL Database')
doc.add_paragraph('Purpose: Persistent data storage with ACID compliance', style='List Bullet')
doc.add_paragraph('Tables:', style='List Bullet')
doc.add_paragraph('  - users: User account information', style='List Bullet')
doc.add_paragraph('  - questions: User questions and answers', style='List Bullet')
doc.add_paragraph('  - documents: Indexed PDF metadata', style='List Bullet')
doc.add_paragraph('  - document_chunks: Text chunks with vector embeddings', style='List Bullet')

doc.add_paragraph('10.4.2 pgvector Extension')
doc.add_paragraph('Purpose: Vector similarity search for semantic retrieval', style='List Bullet')
doc.add_paragraph('Features:', style='List Bullet')
doc.add_paragraph('  - VECTOR data type for storing 384-dimensional embeddings', style='List Bullet')
doc.add_paragraph('  - Cosine distance operator for similarity measurement', style='List Bullet')
doc.add_paragraph('  - IVFFlat index for approximate nearest neighbor search', style='List Bullet')
doc.add_paragraph('Performance:', style='List Bullet')
doc.add_paragraph('  - Sub-millisecond query times for typical workloads', style='List Bullet')
doc.add_paragraph('  - Scalable to millions of vectors with proper indexing', style='List Bullet')

# 10.5 External Services
add_heading_with_bookmark(doc, '10.5 External Services', level=2)

doc.add_paragraph(
    'The application integrates with external AI services for natural language processing '
    'and embedding generation.'
)

doc.add_paragraph('10.5.1 Anthropic Claude API')
doc.add_paragraph('Purpose: Natural language answer generation', style='List Bullet')
doc.add_paragraph('Model: claude-3-haiku-20240307 (optimized for speed and cost)', style='List Bullet')
doc.add_paragraph('Integration:', style='List Bullet')
doc.add_paragraph('  - HTTPS REST API calls via anthropic Python SDK', style='List Bullet')
doc.add_paragraph('  - System prompt configured for Sherlock Holmes domain expertise', style='List Bullet')
doc.add_paragraph('  - Context-aware responses using retrieved document chunks', style='List Bullet')
doc.add_paragraph('Configuration:', style='List Bullet')
doc.add_paragraph('  - API key stored in environment variable (ANTHROPIC_API_KEY)', style='List Bullet')
doc.add_paragraph('  - Max tokens: 1024 per response', style='List Bullet')

doc.add_paragraph('10.5.2 Sentence Transformers')
doc.add_paragraph('Purpose: Text embedding generation for semantic search', style='List Bullet')
doc.add_paragraph('Model: all-MiniLM-L6-v2', style='List Bullet')
doc.add_paragraph('Characteristics:', style='List Bullet')
doc.add_paragraph('  - 384-dimensional output vectors', style='List Bullet')
doc.add_paragraph('  - Local execution (no external API calls)', style='List Bullet')
doc.add_paragraph('  - ~90MB model size, loaded once via Singleton pattern', style='List Bullet')
doc.add_paragraph('Performance:', style='List Bullet')
doc.add_paragraph('  - Fast inference suitable for real-time applications', style='List Bullet')
doc.add_paragraph('  - Batch processing support for document indexing', style='List Bullet')

# 10.6 File Storage
add_heading_with_bookmark(doc, '10.6 File Storage', level=2)

doc.add_paragraph('10.6.1 Books Folder')
doc.add_paragraph('Purpose: Storage location for PDF documents to be indexed', style='List Bullet')
doc.add_paragraph('Location: backend/books/', style='List Bullet')
doc.add_paragraph('Operations:', style='List Bullet')
doc.add_paragraph('  - Read: PDFProcessor extracts text from PDF files', style='List Bullet')
doc.add_paragraph('  - Scan: RAGService scans folder for new documents during indexing', style='List Bullet')
doc.add_paragraph('Security:', style='List Bullet')
doc.add_paragraph('  - Files accessed only by backend service', style='List Bullet')
doc.add_paragraph('  - SHA-256 hash verification prevents duplicate indexing', style='List Bullet')

# 10.7 Data Flow
add_heading_with_bookmark(doc, '10.7 Data Flow', level=2)

doc.add_paragraph('10.7.1 Authentication Flow')
doc.add_paragraph('1. User enters email in LoginPage', style='List Bullet')
doc.add_paragraph('2. authApi.login() sends POST request to /api/auth/login', style='List Bullet')
doc.add_paragraph('3. Auth Routes extract email from request body', style='List Bullet')
doc.add_paragraph('4. AuthService.authenticate_by_email() validates and retrieves user', style='List Bullet')
doc.add_paragraph('5. User Model queries PostgreSQL users table', style='List Bullet')
doc.add_paragraph('6. Response flows back through layers to client', style='List Bullet')
doc.add_paragraph('7. AuthContext stores user state and updates localStorage', style='List Bullet')

doc.add_paragraph('10.7.2 Question Creation Flow')
doc.add_paragraph('1. User fills QuestionForm and clicks Save', style='List Bullet')
doc.add_paragraph('2. questionsApi.create() sends POST to /api/questions', style='List Bullet')
doc.add_paragraph('3. Question Routes validate X-User-ID header and request body', style='List Bullet')
doc.add_paragraph('4. QuestionService.create_question() creates new Question object', style='List Bullet')
doc.add_paragraph('5. Question Model persists to PostgreSQL questions table', style='List Bullet')
doc.add_paragraph('6. Response returns created question to client', style='List Bullet')
doc.add_paragraph('7. HomePage updates state, Toast displays success message', style='List Bullet')

doc.add_paragraph('10.7.3 RAG Answer Generation Flow')
doc.add_paragraph('1. User clicks "Ask Documents" in QuestionForm', style='List Bullet')
doc.add_paragraph('2. ragApi.query() sends POST to /api/rag/query', style='List Bullet')
doc.add_paragraph('3. RAG Routes forward to RAGService.generate_answer()', style='List Bullet')
doc.add_paragraph('4. EmbeddingModel generates query vector (384 dimensions)', style='List Bullet')
doc.add_paragraph('5. DocumentChunk Model performs pgvector similarity search', style='List Bullet')
doc.add_paragraph('6. Top-K relevant chunks retrieved from database', style='List Bullet')
doc.add_paragraph('7. LLMClient sends question + context to Claude API', style='List Bullet')
doc.add_paragraph('8. Generated answer returns through layers to client', style='List Bullet')
doc.add_paragraph('9. QuestionForm populates answer field with response', style='List Bullet')

# 10.8 Deployment Architecture
add_heading_with_bookmark(doc, '10.8 Deployment Architecture', level=2)

doc.add_paragraph('10.8.1 Development Environment')
doc.add_paragraph('Frontend: Vite development server (port 5173)', style='List Bullet')
doc.add_paragraph('Backend: Flask development server (port 5000)', style='List Bullet')
doc.add_paragraph('Database: Local PostgreSQL instance', style='List Bullet')
doc.add_paragraph('Proxy: Vite proxies /api requests to Flask backend', style='List Bullet')

doc.add_paragraph('10.8.2 Production Environment (Docker)')
doc.add_paragraph('Frontend: Nginx serving static React build', style='List Bullet')
doc.add_paragraph('Backend: Gunicorn WSGI server running Flask', style='List Bullet')
doc.add_paragraph('Database: PostgreSQL container with pgvector', style='List Bullet')
doc.add_paragraph('Orchestration: Docker Compose for multi-container management', style='List Bullet')
doc.add_paragraph('Networking: Internal Docker network for service communication', style='List Bullet')

# 10.9 Architecture Diagram
add_heading_with_bookmark(doc, '10.9 Architecture Diagram', level=2)

doc.add_paragraph('File location: diagrams/architecture_diagram.puml')

# Add Architecture Diagram Image
doc.add_picture('diagrams/architecture_diagram.png', width=Inches(6.5))

# Figure Description
fig_para6 = doc.add_paragraph()
fig_para6.alignment = WD_ALIGN_PARAGRAPH.CENTER
fig_run6 = fig_para6.add_run('Figure 6: System Architecture Diagram for Ask Holmes Application')
fig_run6.italic = True

# 10.10 Architecture Decisions
add_heading_with_bookmark(doc, '10.10 Key Architecture Decisions', level=2)

arch_decision_table = doc.add_table(rows=1, cols=3)
arch_decision_table.style = 'Table Grid'
hdr = arch_decision_table.rows[0].cells
hdr[0].text = 'Decision'
hdr[1].text = 'Rationale'
hdr[2].text = 'Trade-offs'
for cell in hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

decisions = [
    ('React SPA Frontend', 'Rich interactivity, component reusability, large ecosystem', 'Initial load time, SEO challenges (mitigated for internal app)'),
    ('Flask Backend', 'Lightweight, Python ecosystem for ML/AI, rapid development', 'Less suited for high-concurrency (mitigated with Gunicorn)'),
    ('PostgreSQL + pgvector', 'ACID compliance, vector search in single database, mature ecosystem', 'Requires PostgreSQL-specific deployment'),
    ('Local Embedding Model', 'No API costs, low latency, data privacy', 'Memory usage (~90MB), CPU-bound processing'),
    ('Claude API for LLM', 'High-quality responses, managed service, cost-effective', 'External dependency, API rate limits, network latency'),
    ('REST API Design', 'Stateless, cacheable, widely understood', 'Multiple requests for complex operations'),
    ('Service Layer Pattern', 'Testability, separation of concerns, reusability', 'Additional abstraction layer'),
    ('Docker Deployment', 'Consistent environments, easy scaling, isolation', 'Container orchestration complexity'),
]

for decision, rationale, tradeoffs in decisions:
    row = arch_decision_table.add_row().cells
    row[0].text = decision
    row[1].text = rationale
    row[2].text = tradeoffs

# ============================================
# 11. CONCLUSION
# ============================================
add_heading_with_bookmark(doc, '11. Conclusion', level=1)

# 11.1 Summary
add_heading_with_bookmark(doc, '11.1 Summary', level=2)
doc.add_paragraph(
    'This Software Requirements Specification document has provided a comprehensive overview of the '
    'Ask Holmes application, a Retrieval-Augmented Generation (RAG) system designed to serve as an '
    'intelligent knowledge base for Sherlock Holmes literature. The document has covered all aspects '
    'of the system from functional and non-functional requirements to detailed architectural design '
    'and implementation specifications.'
)

doc.add_paragraph('Key highlights of this specification include:')
doc.add_paragraph(
    'Comprehensive Requirements: The document defines 45+ functional requirements covering user '
    'authentication, question management, answer generation, RAG pipeline operations, and document '
    'management capabilities.',
    style='List Bullet'
)
doc.add_paragraph(
    'Robust Architecture: A three-tier architecture (Client, Server, Data) with clear separation '
    'of concerns, leveraging React for the frontend, Flask for the backend, and PostgreSQL with '
    'pgvector for semantic search capabilities.',
    style='List Bullet'
)
doc.add_paragraph(
    'AI-Powered Features: Integration of state-of-the-art AI technologies including Sentence '
    'Transformers for embedding generation and Anthropic Claude API for natural language response '
    'generation.',
    style='List Bullet'
)
doc.add_paragraph(
    'Quality Assurance: Comprehensive unit test cases covering all layers of the application, '
    'ensuring reliability and maintainability of the codebase.',
    style='List Bullet'
)
doc.add_paragraph(
    'Scalable Design: Docker-based deployment architecture enabling consistent environments '
    'and easy horizontal scaling.',
    style='List Bullet'
)

# 11.2 Project Scope Achievement
add_heading_with_bookmark(doc, '11.2 Project Scope Achievement', level=2)
doc.add_paragraph(
    'The Ask Holmes application successfully addresses the core objectives outlined in the project scope:'
)

scope_table = doc.add_table(rows=1, cols=3)
scope_table.style = 'Table Grid'
scope_hdr = scope_table.rows[0].cells
scope_hdr[0].text = 'Objective'
scope_hdr[1].text = 'Implementation'
scope_hdr[2].text = 'Status'
for cell in scope_hdr:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True

scope_items = [
    ('Question Management', 'Full CRUD operations with user isolation', 'Achieved'),
    ('AI-Powered Answers', 'RAG pipeline with semantic search and LLM generation', 'Achieved'),
    ('Trusted Sources', 'Document indexing with source attribution', 'Achieved'),
    ('User Authentication', 'Email-based authentication with session management', 'Achieved'),
    ('Responsive UI', 'React-based SPA with modern component architecture', 'Achieved'),
    ('Data Persistence', 'PostgreSQL database with pgvector extension', 'Achieved'),
    ('Containerized Deployment', 'Docker Compose multi-container setup', 'Achieved'),
]

for objective, implementation, status in scope_items:
    row = scope_table.add_row().cells
    row[0].text = objective
    row[1].text = implementation
    row[2].text = status

# 11.3 Future Enhancements
add_heading_with_bookmark(doc, '11.3 Future Enhancements', level=2)
doc.add_paragraph(
    'While the current implementation fulfills all specified requirements, the following '
    'enhancements could be considered for future versions:'
)

doc.add_paragraph(
    'Advanced Authentication: Implementation of OAuth 2.0, multi-factor authentication, '
    'and password-based login for enhanced security.',
    style='List Bullet'
)
doc.add_paragraph(
    'Extended Document Support: Support for additional document formats (EPUB, TXT, HTML) '
    'and multi-language content indexing.',
    style='List Bullet'
)
doc.add_paragraph(
    'Enhanced RAG Capabilities: Implementation of hybrid search (combining keyword and semantic), '
    'query rewriting, and multi-hop reasoning for complex questions.',
    style='List Bullet'
)
doc.add_paragraph(
    'Analytics Dashboard: User analytics, question trends, and system usage metrics '
    'for administrators.',
    style='List Bullet'
)
doc.add_paragraph(
    'Collaborative Features: Shared question collections, user annotations, and '
    'collaborative knowledge building.',
    style='List Bullet'
)
doc.add_paragraph(
    'Mobile Application: Native mobile applications for iOS and Android platforms.',
    style='List Bullet'
)

# 11.4 Technical Debt and Recommendations
add_heading_with_bookmark(doc, '11.4 Technical Debt and Recommendations', level=2)
doc.add_paragraph(
    'The following technical recommendations should be addressed to maintain and improve '
    'the system over time:'
)

doc.add_paragraph(
    'API Versioning: Implement API versioning (e.g., /api/v1/) to support backward '
    'compatibility during future updates.',
    style='List Bullet'
)
doc.add_paragraph(
    'Caching Layer: Introduce Redis caching for frequently accessed data and embedding '
    'results to improve response times.',
    style='List Bullet'
)
doc.add_paragraph(
    'Monitoring and Logging: Implement comprehensive logging with ELK stack or similar '
    'solutions for production monitoring.',
    style='List Bullet'
)
doc.add_paragraph(
    'CI/CD Pipeline: Establish automated testing and deployment pipelines for '
    'continuous integration and delivery.',
    style='List Bullet'
)
doc.add_paragraph(
    'Load Testing: Conduct thorough load testing to identify performance bottlenecks '
    'under high concurrency scenarios.',
    style='List Bullet'
)

# 11.5 Final Remarks
add_heading_with_bookmark(doc, '11.5 Final Remarks', level=2)
doc.add_paragraph(
    'The Ask Holmes application represents a modern implementation of RAG technology, '
    'combining the rich literary world of Sherlock Holmes with cutting-edge AI capabilities. '
    'The system architecture ensures maintainability, scalability, and extensibility, '
    'providing a solid foundation for future development.'
)

doc.add_paragraph(
    'This SRS document serves as the authoritative reference for the Ask Holmes application, '
    'providing all necessary information for development, testing, deployment, and maintenance. '
    'All stakeholders are encouraged to refer to this document throughout the project lifecycle '
    'and to propose updates as the system evolves.'
)

doc.add_paragraph(
    'The successful implementation of this specification will deliver a valuable tool for '
    'Sherlock Holmes enthusiasts, researchers, and casual readers alike, enabling them to '
    'explore and interact with the classic detective stories in an innovative and engaging way.'
)

# Save the document
doc.save('SRS_Ask_Holmes.docx')

print("SRS document created successfully: SRS_Ask_Holmes.docx")
