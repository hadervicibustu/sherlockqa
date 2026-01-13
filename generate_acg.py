from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_heading_with_style(doc, text, level):
    """Add a heading with consistent styling."""
    heading = doc.add_heading(text, level=level)
    return heading

def create_table_with_header(doc, headers, col_widths=None):
    """Create a table with styled header row."""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    header_row = table.rows[0]
    for i, header_text in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header_text
        set_cell_shading(cell, '2E4057')  # Dark blue
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.size = Pt(10)

    # Set column widths if provided
    if col_widths:
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = Inches(width)

    return table

def add_table_row(table, cells_data, alternate=False):
    """Add a row to the table with optional alternating color."""
    row = table.add_row()
    for i, text in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = text
        if alternate:
            set_cell_shading(cell, 'F0F4F8')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(9)
    return row

# Create a new Document
doc = Document()

# ============================================================================
# TITLE PAGE
# ============================================================================
title = doc.add_heading('Agentic Coding Guidelines', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Claude Code Agent Configuration Documentation')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in subtitle.runs:
    run.font.size = Pt(14)
    run.font.italic = True

# Add project info
project_para = doc.add_paragraph()
project_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
project_run = project_para.add_run('Project: Sherlock QA (Ask Holmes)')
project_run.font.size = Pt(12)

version = doc.add_paragraph('Version 1.0')
version.alignment = WD_ALIGN_PARAGRAPH.CENTER

date_para = doc.add_paragraph('January 2026')
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add section break
doc.add_section()

# Set vertical alignment for title page
first_section = doc.sections[0]
sectPr = first_section._sectPr
vAlign = OxmlElement('w:vAlign')
vAlign.set(qn('w:val'), 'center')
sectPr.append(vAlign)

# Page numbers
second_section = doc.sections[1]
second_section.footer.is_linked_to_previous = False
footer = second_section.footer
footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer_para.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'separate')
fldChar3 = OxmlElement('w:fldChar')
fldChar3.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)
run._r.append(fldChar3)

# ============================================================================
# TABLE OF CONTENTS
# ============================================================================
add_heading_with_style(doc, 'Table of Contents', 1)

toc_items = [
    '1. Executive Summary',
    '2. Integration Summary Table',
    '3. MCP Servers',
    '4. Custom Skills',
    '5. Available Tools',
    '6. Sub-Agents',
    '7. Configuration Reference'
]

for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(6)

doc.add_page_break()

# ============================================================================
# 1. EXECUTIVE SUMMARY
# ============================================================================
add_heading_with_style(doc, '1. Executive Summary', 1)

summary_text = """This document provides comprehensive documentation for the Claude Code agent configuration used in the Sherlock QA project. It catalogs all custom integrations, MCP servers, skills, tools, and sub-agents configured for this specific development environment.

The agent is powered by Claude Opus 4.5 (model ID: claude-opus-4-5-20251101) and operates within a Windows development environment with access to browser automation, file operations, web services, and specialized coding assistance capabilities."""

p = doc.add_paragraph(summary_text)
p.paragraph_format.space_after = Pt(12)

add_heading_with_style(doc, '1.1 Scope', 2)
scope_text = """This documentation covers:
- MCP (Model Context Protocol) server connections
- Custom skills for project-specific operations
- Complete tool catalog organized by category
- Sub-agent configurations for delegated tasks

Excluded: General Claude capabilities, standard LLM features, and verbose explanations of common functionality."""
doc.add_paragraph(scope_text)

doc.add_page_break()

# ============================================================================
# 2. INTEGRATION SUMMARY TABLE
# ============================================================================
add_heading_with_style(doc, '2. Integration Summary Table', 1)

p = doc.add_paragraph('Overview of all configured integrations in this Claude agent setup:')
p.paragraph_format.space_after = Pt(12)

# Summary table
summary_table = create_table_with_header(doc, ['Category', 'Component', 'Count', 'Status'], [1.5, 2.5, 1, 1])

summary_data = [
    ('MCP Servers', 'Playwright Browser Automation', '1', 'Active'),
    ('Custom Skills', 'register-user, index-books', '2', 'Active'),
    ('Core Tools', 'File, Search, Edit, Write', '8', 'Active'),
    ('Planning Tools', 'EnterPlanMode, ExitPlanMode, TodoWrite', '3', 'Active'),
    ('Communication', 'AskUserQuestion', '1', 'Active'),
    ('Web Tools', 'WebFetch, WebSearch', '2', 'Active'),
    ('Browser Tools', 'Playwright MCP Tools', '22', 'Active'),
    ('Sub-Agents', 'Specialized Task Agents', '6', 'Active'),
]

for i, row_data in enumerate(summary_data):
    add_table_row(summary_table, row_data, alternate=(i % 2 == 1))

doc.add_paragraph()  # Spacing

# Totals
totals_para = doc.add_paragraph()
totals_run = totals_para.add_run('Total Integrations: 45 components across 8 categories')
totals_run.font.bold = True

doc.add_page_break()

# ============================================================================
# 3. MCP SERVERS
# ============================================================================
add_heading_with_style(doc, '3. MCP Servers', 1)

intro = doc.add_paragraph('Model Context Protocol (MCP) servers extend Claude\'s capabilities by providing specialized tool sets for specific domains.')
intro.paragraph_format.space_after = Pt(12)

add_heading_with_style(doc, '3.1 Playwright MCP Server', 2)

playwright_desc = """The Playwright MCP server provides comprehensive browser automation capabilities for web testing, scraping, and interaction tasks."""
doc.add_paragraph(playwright_desc)

# Playwright tools table
add_heading_with_style(doc, 'Playwright Tools Reference', 3)

playwright_table = create_table_with_header(doc, ['Tool Name', 'Function', 'Key Parameters'], [2.2, 2.5, 2])

playwright_tools = [
    ('browser_navigate', 'Navigate to a URL', 'url (required)'),
    ('browser_navigate_back', 'Go back to previous page', 'None'),
    ('browser_snapshot', 'Capture accessibility snapshot', 'filename (optional)'),
    ('browser_take_screenshot', 'Take page screenshot', 'filename, type, fullPage'),
    ('browser_click', 'Click on element', 'element, ref, button'),
    ('browser_type', 'Type text into element', 'element, ref, text, submit'),
    ('browser_fill_form', 'Fill multiple form fields', 'fields[]'),
    ('browser_select_option', 'Select dropdown option', 'element, ref, values[]'),
    ('browser_hover', 'Hover over element', 'element, ref'),
    ('browser_drag', 'Drag and drop', 'startRef, endRef'),
    ('browser_press_key', 'Press keyboard key', 'key'),
    ('browser_file_upload', 'Upload files', 'paths[]'),
    ('browser_evaluate', 'Execute JavaScript', 'function, element'),
    ('browser_run_code', 'Run Playwright code', 'code'),
    ('browser_tabs', 'Manage browser tabs', 'action (list/new/close/select)'),
    ('browser_wait_for', 'Wait for condition', 'text, textGone, time'),
    ('browser_handle_dialog', 'Handle browser dialogs', 'accept, promptText'),
    ('browser_console_messages', 'Get console messages', 'level'),
    ('browser_network_requests', 'Get network requests', 'includeStatic'),
    ('browser_resize', 'Resize browser window', 'width, height'),
    ('browser_close', 'Close the page', 'None'),
    ('browser_install', 'Install browser', 'None'),
]

for i, row_data in enumerate(playwright_tools):
    add_table_row(playwright_table, row_data, alternate=(i % 2 == 1))

doc.add_page_break()

# ============================================================================
# 4. CUSTOM SKILLS
# ============================================================================
add_heading_with_style(doc, '4. Custom Skills', 1)

skills_intro = """Custom skills are project-specific capabilities configured for the Sherlock QA application. These skills provide shortcuts for common operations and can be invoked using the /skill-name syntax."""
doc.add_paragraph(skills_intro)

skills_table = create_table_with_header(doc, ['Skill Name', 'Display Name', 'Description', 'Invocation'], [1.5, 1.5, 2.5, 1.2])

skills_data = [
    ('register-user', 'Register User', 'Registers a new user in the Sherlock QA system for authentication and access control', '/register-user'),
    ('index-books', 'Index Books for RAG', 'Indexes book content for Retrieval-Augmented Generation, enabling semantic search across the document corpus', '/index-books'),
]

for i, row_data in enumerate(skills_data):
    add_table_row(skills_table, row_data, alternate=(i % 2 == 1))

doc.add_paragraph()

add_heading_with_style(doc, '4.1 Skill Usage', 2)
usage_text = """Skills are invoked through the Skill tool with the following pattern:
- Direct invocation: /skill-name (e.g., /register-user)
- With arguments: /skill-name args (e.g., /index-books --path /data)

Note: Skills are loaded from /mnt/skills/ or equivalent configuration directory."""
doc.add_paragraph(usage_text)

doc.add_page_break()

# ============================================================================
# 5. AVAILABLE TOOLS
# ============================================================================
add_heading_with_style(doc, '5. Available Tools', 1)

tools_intro = """This section catalogs all tools available to the Claude agent, organized by functional category."""
doc.add_paragraph(tools_intro)

# 5.1 Core File Operations
add_heading_with_style(doc, '5.1 Core File Operations', 2)

file_tools_table = create_table_with_header(doc, ['Tool', 'Purpose', 'Key Capabilities'], [1.5, 2, 3.2])

file_tools = [
    ('Read', 'Read file contents', 'Absolute paths, line offset/limit, supports images, PDFs, notebooks'),
    ('Write', 'Create/overwrite files', 'Full file creation, requires prior Read for existing files'),
    ('Edit', 'Modify existing files', 'String replacement, replace_all option, preserves indentation'),
    ('Glob', 'Pattern-based file search', 'Glob patterns (e.g., **/*.py), sorted by modification time'),
    ('Grep', 'Content search (ripgrep)', 'Regex patterns, file type filters, context lines (-A/-B/-C)'),
    ('NotebookEdit', 'Edit Jupyter notebooks', 'Cell replacement, insert, delete operations'),
]

for i, row_data in enumerate(file_tools):
    add_table_row(file_tools_table, row_data, alternate=(i % 2 == 1))

# 5.2 Execution Tools
add_heading_with_style(doc, '5.2 Execution Tools', 2)

exec_tools_table = create_table_with_header(doc, ['Tool', 'Purpose', 'Key Capabilities'], [1.5, 2, 3.2])

exec_tools = [
    ('Bash', 'Execute shell commands', 'Persistent session, timeout support, background execution'),
    ('Task', 'Launch specialized agents', 'Sub-agent delegation, background tasks, context sharing'),
    ('TaskOutput', 'Retrieve task results', 'Blocking/non-blocking, timeout configuration'),
    ('KillShell', 'Terminate background shell', 'Shell ID-based termination'),
]

for i, row_data in enumerate(exec_tools):
    add_table_row(exec_tools_table, row_data, alternate=(i % 2 == 1))

# 5.3 Planning & Organization
add_heading_with_style(doc, '5.3 Planning & Organization', 2)

planning_tools_table = create_table_with_header(doc, ['Tool', 'Purpose', 'Key Capabilities'], [1.5, 2, 3.2])

planning_tools = [
    ('TodoWrite', 'Task management', 'Create/update task lists, status tracking (pending/in_progress/completed)'),
    ('EnterPlanMode', 'Start planning mode', 'Transitions to plan mode for complex implementations'),
    ('ExitPlanMode', 'Complete planning', 'Signals plan completion, requests user approval'),
]

for i, row_data in enumerate(planning_tools):
    add_table_row(planning_tools_table, row_data, alternate=(i % 2 == 1))

# 5.4 Communication & Web
add_heading_with_style(doc, '5.4 Communication & Web Tools', 2)

comm_tools_table = create_table_with_header(doc, ['Tool', 'Purpose', 'Key Capabilities'], [1.5, 2, 3.2])

comm_tools = [
    ('AskUserQuestion', 'Interactive queries', 'Multi-choice questions, multi-select support, 2-4 options'),
    ('WebFetch', 'Fetch URL content', 'HTML to markdown conversion, AI processing, redirect handling'),
    ('WebSearch', 'Web search', 'Real-time search, domain filtering, source citations required'),
    ('Skill', 'Execute custom skills', 'Invoke configured skills, argument passing'),
]

for i, row_data in enumerate(comm_tools):
    add_table_row(comm_tools_table, row_data, alternate=(i % 2 == 1))

doc.add_page_break()

# ============================================================================
# 6. SUB-AGENTS
# ============================================================================
add_heading_with_style(doc, '6. Sub-Agents', 1)

subagent_intro = """Sub-agents are specialized agents launched via the Task tool for handling complex, multi-step tasks autonomously. Each agent type has specific capabilities and tool access."""
doc.add_paragraph(subagent_intro)

subagent_table = create_table_with_header(doc, ['Agent Type', 'Specialization', 'Available Tools', 'Use Case'], [1.4, 1.8, 1.5, 2])

subagents = [
    ('Bash', 'Command execution', 'Bash', 'Git operations, terminal tasks, system commands'),
    ('general-purpose', 'Multi-step tasks', 'All tools (*)', 'Research, code search, complex workflows'),
    ('Explore', 'Codebase exploration', 'All tools', 'File pattern search, code analysis, architecture review'),
    ('Plan', 'Implementation planning', 'All tools', 'Design strategies, architectural decisions, step-by-step plans'),
    ('claude-code-guide', 'Claude Code expertise', 'Glob, Grep, Read, WebFetch, WebSearch', 'CLI features, MCP servers, API usage questions'),
    ('statusline-setup', 'Status line config', 'Read, Edit', 'Configure Claude Code status line settings'),
]

for i, row_data in enumerate(subagents):
    add_table_row(subagent_table, row_data, alternate=(i % 2 == 1))

doc.add_paragraph()

add_heading_with_style(doc, '6.1 Sub-Agent Invocation', 2)

invocation_text = """Sub-agents are invoked through the Task tool with the following parameters:
- subagent_type: Required. Specifies the agent type (e.g., "Explore", "Plan")
- prompt: Required. The task description for the agent
- description: Required. Short 3-5 word summary
- model: Optional. Override model (sonnet, opus, haiku)
- run_in_background: Optional. Execute asynchronously
- resume: Optional. Continue from previous agent session"""
doc.add_paragraph(invocation_text)

add_heading_with_style(doc, '6.2 Agent Selection Guidelines', 2)

guidelines = doc.add_paragraph()
guidelines_text = """- Use Explore for codebase navigation and understanding
- Use Plan for complex features requiring architectural decisions
- Use Bash for git operations and command sequences
- Use general-purpose for multi-tool research tasks
- Use claude-code-guide for Claude Code CLI/API questions"""
doc.add_paragraph(guidelines_text)

doc.add_page_break()

# ============================================================================
# 7. CONFIGURATION REFERENCE
# ============================================================================
add_heading_with_style(doc, '7. Configuration Reference', 1)

add_heading_with_style(doc, '7.1 Environment Details', 2)

env_table = create_table_with_header(doc, ['Property', 'Value'], [2.5, 4.2])

env_data = [
    ('Platform', 'Windows (win32)'),
    ('Working Directory', 'C:\\Users\\hamid\\OneDrive\\Desktop\\sherlock'),
    ('Git Repository', 'Yes'),
    ('Current Branch', 'feature/add-srs-doc'),
    ('Model', 'Claude Opus 4.5 (claude-opus-4-5-20251101)'),
    ('Knowledge Cutoff', 'May 2025'),
]

for i, row_data in enumerate(env_data):
    add_table_row(env_table, row_data, alternate=(i % 2 == 1))

add_heading_with_style(doc, '7.2 Tool Usage Best Practices', 2)

best_practices = """1. File Operations: Always use Read before Edit/Write for existing files
2. Search: Use Glob for file patterns, Grep for content search
3. Exploration: Prefer Task tool with Explore agent for open-ended searches
4. Planning: Use EnterPlanMode for complex multi-file changes
5. Task Tracking: Use TodoWrite proactively for multi-step tasks
6. Browser Automation: Use browser_snapshot over screenshots for interactions
7. Git Operations: Never force push, always verify before committing"""
doc.add_paragraph(best_practices)

add_heading_with_style(doc, '7.3 Security Guidelines', 2)

security_text = """- Never commit files containing secrets (.env, credentials.json)
- Avoid command injection vulnerabilities in Bash commands
- Validate external input at system boundaries
- Use authorization context for security testing tools
- Refuse requests for malicious code or destructive operations"""
doc.add_paragraph(security_text)

# ============================================================================
# SAVE DOCUMENT
# ============================================================================
output_path = 'Agentic_Coding_Guidelines.docx'
doc.save(output_path)
print(f"Document saved to: {output_path}")
